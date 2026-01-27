from fastapi import APIRouter, Request, HTTPException,Query
from fastapi.responses import FileResponse
from states.global_state import influx_state, redis_state, os_spec
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from typing import Dict, Any, List, Optional
import numpy as np
import pandas as pd
import pyarrow.parquet as pq
from datetime import datetime, timedelta, timezone
import logging, os, json, tempfile, uuid
from utils.en50160_dataMap import EN50160ReportProcessor, WeeklyReportConfig
from .api import get_asset, get_params, get_trendData, Trend
from utils.weekly_report import generate_weekly_report
from pathlib import Path
import warnings

router = APIRouter()

REPORTS_DIR = Path("/usr/local/sv500/reports")

# matplotlib 기본 설정
plt.rcParams['axes.unicode_minus'] = False

# matplotlib 폰트 관련 경고 억제
warnings.filterwarnings('ignore', category=UserWarning, module='matplotlib')
warnings.filterwarnings('ignore', message='.*glyph.*')
warnings.filterwarnings('ignore', message='.*Substituting.*')

# matplotlib 로거 레벨 조정
logging.getLogger('matplotlib').setLevel(logging.DEBUG)
logging.getLogger('matplotlib.font_manager').setLevel(logging.DEBUG)

_font_prop_cache = None
_font_initialized = False

# ============================================
# 한글 폰트 설정
# ============================================
def setup_korean_font():
    """한글 폰트 설정 (한 번만 초기화)"""
    global _font_prop_cache, _font_initialized

    if _font_initialized:
        return _font_prop_cache

    font_path = '/home/root/webserver/fonts/NanumGothicCoding.ttf'
    if os.path.exists(font_path):
        try:
            fm.fontManager.addfont(font_path)
            font_prop = fm.FontProperties(fname=font_path)
            font_name = font_prop.get_name()
            plt.rcParams['font.family'] = font_name
            plt.rcParams['axes.unicode_minus'] = False
            _font_prop_cache = font_prop
            _font_initialized = True
            return font_prop
        except Exception as e:
            logging.warning(f"폰트 로드 실패: {e}")

    plt.rcParams['axes.unicode_minus'] = False
    _font_initialized = True
    return None


# 모듈 로드 시 폰트 초기 설정
setup_korean_font()

processor = EN50160ReportProcessor()

# ============================================
# 유틸리티 함수
# ============================================
def format_xaxis_label(label_str):
    """
    X축 라벨 포맷팅
    "2025-12-11T09:30:00+09:00" → "12-11 09:30"
    """
    try:
        if isinstance(label_str, str):
            if 'T' in label_str:
                date_part = label_str.split('T')[0]
                time_part = label_str.split('T')[1][:5]
                month_day = date_part[5:]
                return f"{month_day} {time_part}"
        return str(label_str)[-8:]
    except:
        return str(label_str)[-8:]


def parse_asset_data(api_response):
    """
    API 응답을 리포트용 데이터로 변환
    """
    if not api_response.get('success'):
        return None

    data_list = api_response.get('data', [])
    mac = api_response.get('mac', '-')
    drive_type = api_response.get('driveType', '')

    parsed = {
        'mac': mac,
        'drive_type': drive_type,
        'rated_voltage': '-',
        'rated_current': '-',
        'rated_power': '-',
        'rated_capacity': '-',
        'frequency': 60,
        'pt_wiring_mode': '-'
    }

    for item in data_list:
        name = item.get('Name', '')
        value = item.get('Value', '-')

        if name == 'RatedVoltage':
            parsed['rated_voltage'] = value
        elif name == 'RatedCurrent':
            parsed['rated_current'] = value
        elif name == 'RatedFrequency':
            parsed['frequency'] = value
        elif name == 'RatedKVA':
            parsed['rated_capacity'] = value
        elif name == 'PT_WiringMode':
            parsed['pt_wiring_mode'] = value

    return parsed


# ============================================
# 설비 정보 섹션 생성
# ============================================
def create_channel_info_section(doc, channel_data, locale='en'):
    """
    Channel Information 섹션 생성
    """
    # === 헤더 ===
    header_table = doc.add_table(rows=1, cols=1)
    header_table.style = 'Table Grid'

    title_cell = header_table.rows[0].cells[0]
    title_para = title_cell.paragraphs[0]

    if locale == 'ko':
        title_text = f'{channel_data.get("channel", "Main")} 채널 정보'
    else:
        title_text = f'{channel_data.get("channel", "Main")} Channel Information'

    title_run = title_para.add_run(title_text)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '64748B')
    title_cell._element.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()

    # === 첫 번째 줄: 이름/타입/구동방식 ===
    info_table1 = doc.add_table(rows=1, cols=3)
    info_table1.style = 'Light Grid Accent 1'

    cells = info_table1.rows[0].cells

    # 이름
    name_cell = cells[0]
    name_label = name_cell.paragraphs[0]
    label_text = '설비명' if locale == 'ko' else 'Asset Name'
    name_label_run = name_label.add_run(f'{label_text}\n')
    name_label_run.font.size = Pt(9)
    name_label_run.font.bold = True
    name_label_run.font.color.rgb = RGBColor(107, 114, 128)

    name_value = name_cell.add_paragraph()
    name_value_run = name_value.add_run(channel_data.get('name', '-'))
    name_value_run.font.size = Pt(14)
    name_value_run.font.bold = True
    name_value_run.font.color.rgb = RGBColor(31, 41, 55)

    # 타입
    type_cell = cells[1]
    type_label = type_cell.paragraphs[0]
    label_text = '설비 타입' if locale == 'ko' else 'Asset Type'
    type_label_run = type_label.add_run(f'{label_text}\n')
    type_label_run.font.size = Pt(9)
    type_label_run.font.bold = True
    type_label_run.font.color.rgb = RGBColor(107, 114, 128)

    type_value = type_cell.add_paragraph()
    type_value_run = type_value.add_run(channel_data.get('type', '-'))
    type_value_run.font.size = Pt(14)
    type_value_run.font.bold = True
    type_value_run.font.color.rgb = RGBColor(31, 41, 55)

    # Drive Type (Motor인 경우만)
    asset_type = channel_data.get('type', '').lower()
    if asset_type in ['motor', 'motorfeed', 'pump', 'fan', 'compressor']:
        drive_cell = cells[2]
        drive_label = drive_cell.paragraphs[0]
        label_text = '구동 방식' if locale == 'ko' else 'Drive Type'
        drive_label_run = drive_label.add_run(f'{label_text}\n')
        drive_label_run.font.size = Pt(9)
        drive_label_run.font.bold = True
        drive_label_run.font.color.rgb = RGBColor(107, 114, 128)

        drive_value = drive_cell.add_paragraph()
        drive_type = channel_data.get('drive_type', '')
        if locale == 'ko':
            drive_type_text = '직입 기동' if drive_type == 'DOL' else '인버터' if drive_type == 'VFD' else '-'
        else:
            drive_type_text = 'DOL' if drive_type == 'DOL' else 'VFD' if drive_type == 'VFD' else '-'
        drive_value_run = drive_value.add_run(drive_type_text)
        drive_value_run.font.size = Pt(14)
        drive_value_run.font.bold = True
        drive_value_run.font.color.rgb = RGBColor(31, 41, 55)

    doc.add_paragraph()

    # === 두 번째 줄: 상세 사양 ===
    spec_items = []

    if locale == 'ko':
        if 'transformer' in asset_type:
            spec_items = [
                {'name': '정격용량', 'value': channel_data.get('rated_capacity', '-'), 'unit': 'kVA'},
                {'name': '정격전압', 'value': channel_data.get('rated_voltage', '-'), 'unit': 'V'},
                {'name': '정격주파수', 'value': channel_data.get('frequency', 60), 'unit': 'Hz'},
                {'name': 'PT 결선방식', 'value': channel_data.get('pt_wiring_mode', '-'), 'unit': ''},
                {'name': '위치', 'value': channel_data.get('location', '-'), 'unit': ''},
                {'name': 'MAC', 'value': channel_data.get('mac', '-'), 'unit': ''},
            ]
        else:
            spec_items = [
                {'name': '정격전압', 'value': channel_data.get('rated_voltage', '-'), 'unit': 'V'},
                {'name': '정격전류', 'value': channel_data.get('rated_current', '-'), 'unit': 'A'},
                {'name': 'PT 결선방식', 'value': channel_data.get('pt_wiring_mode', '-'), 'unit': ''},
                {'name': '정격주파수', 'value': channel_data.get('frequency', 60), 'unit': 'Hz'},
                {'name': '위치', 'value': channel_data.get('location', '-'), 'unit': ''},
                {'name': 'MAC', 'value': channel_data.get('mac', '-'), 'unit': ''},
            ]
    else:
        if 'transformer' in asset_type:
            spec_items = [
                {'name': 'Rated Capacity', 'value': channel_data.get('rated_capacity', '-'), 'unit': 'kVA'},
                {'name': 'Rated Voltage', 'value': channel_data.get('rated_voltage', '-'), 'unit': 'V'},
                {'name': 'Rated Frequency', 'value': channel_data.get('frequency', 60), 'unit': 'Hz'},
                {'name': 'PT Wiring', 'value': channel_data.get('pt_wiring_mode', '-'), 'unit': ''},
                {'name': 'Location', 'value': channel_data.get('location', '-'), 'unit': ''},
                {'name': 'MAC', 'value': channel_data.get('mac', '-'), 'unit': ''},
            ]
        else:
            spec_items = [
                {'name': 'Rated Voltage', 'value': channel_data.get('rated_voltage', '-'), 'unit': 'V'},
                {'name': 'Rated Current', 'value': channel_data.get('rated_current', '-'), 'unit': 'A'},
                {'name': 'PT Wiring', 'value': channel_data.get('pt_wiring_mode', '-'), 'unit': ''},
                {'name': 'Rated Frequency', 'value': channel_data.get('frequency', 60), 'unit': 'Hz'},
                {'name': 'Location', 'value': channel_data.get('location', '-'), 'unit': ''},
                {'name': 'MAC', 'value': channel_data.get('mac', '-'), 'unit': ''},
            ]

    num_cols = 3
    num_rows = (len(spec_items) + num_cols - 1) // num_cols

    spec_table = doc.add_table(rows=num_rows, cols=num_cols)
    spec_table.style = 'Light Grid Accent 1'

    for idx, item in enumerate(spec_items):
        row_idx = idx // num_cols
        col_idx = idx % num_cols

        cell = spec_table.rows[row_idx].cells[col_idx]

        label_para = cell.paragraphs[0]
        label_run = label_para.add_run(f'{item["name"]}\n')
        label_run.font.size = Pt(9)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(107, 114, 128)

        value_para = cell.add_paragraph()
        value_text = f'{item["value"]} {item["unit"]}'.strip()
        value_run = value_para.add_run(value_text)
        value_run.font.size = Pt(12)
        value_run.font.bold = True
        value_run.font.color.rgb = RGBColor(31, 41, 55)


# ============================================
# 진단 바차트 생성
# ============================================
def create_diagnosis_bar_chart(chart_path, main_data, mode='diagnosis', locale='en'):
    """
    진단 결과 바차트 생성 (matplotlib)
    """
    # 항상 한글 폰트 설정
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(12, 6))

    labels = []
    values = []
    colors = []

    color_map = {
        0: '#64748B',
        1: '#10B981',
        2: '#EAB308',
        3: '#F97316',
        4: '#EF4444',
    }

    for item in main_data:
        if locale == 'ko':
            label = item.get('title_ko') or item.get('title') or item.get('item_name', '')
        elif locale == 'ja':
            label = item.get('title_ja') or item.get('title') or item.get('item_name', '')
        else:
            label = item.get('title_en') or item.get('title') or item.get('item_name', '')
        labels.append(label)

        status = item.get('status', 0)
        values.append(status if status > 0 else 0.2)
        colors.append(color_map.get(int(status), color_map[0]))

    y_pos = np.arange(len(labels))
    ax.barh(y_pos, values, color=colors, height=0.6, edgecolor='white', linewidth=1)

    ax.set_yticks(y_pos)
    # 폰트 직접 적용
    if font_prop:
        ax.set_yticklabels(labels, fontsize=11, fontproperties=font_prop)
    else:
        ax.set_yticklabels(labels, fontsize=11)
    ax.set_xlim(0, 4.5)
    ax.set_xticks([0, 1, 2, 3, 4])

    if mode == 'powerquality':
        if locale == 'ko':
            status_labels = ['정지', '정상', '낮음', '중간', '높음']
        else:
            status_labels = ['Stop', 'OK', 'Low', 'Medium', 'High']
    else:
        if locale == 'ko':
            status_labels = ['정지', '정상', '주의', '경고', '위험']
        else:
            status_labels = ['Stop', 'OK', 'Warning', 'Inspect', 'Repair']

    # 폰트 직접 적용
    if font_prop:
        ax.set_xticklabels(status_labels, fontsize=10, fontproperties=font_prop)
    else:
        ax.set_xticklabels(status_labels, fontsize=10)
    ax.xaxis.grid(True, linestyle='--', alpha=0.3)
    ax.set_axisbelow(True)

    if mode == 'powerquality':
        if locale == 'ko':
            title = '전력품질 진단 현황'
        else:
            title = 'Power Quality Status'
    else:
        if locale == 'ko':
            title = '설비 진단 현황'
        else:
            title = 'Diagnostic Status'

    # 폰트 직접 적용
    if font_prop:
        ax.set_title(title, fontsize=14, fontweight='bold', pad=15, fontproperties=font_prop)
    else:
        ax.set_title(title, fontsize=14, fontweight='bold', pad=15)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


# ============================================
# 트렌드 차트 생성
# ============================================
def create_trend_chart(chart_path, trend_data, locale='en'):
    """
    트렌드 차트 생성 (matplotlib)
    """
    # 항상 한글 폰트 설정
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(10, 4))

    labels = trend_data.get('lineLabels', [])
    datasets = trend_data.get('lineData', [])
    chart_title = trend_data.get('lineTitle', 'Trend')

    if not labels or not datasets:
        plt.close()
        return None

    x = np.arange(len(labels))

    for dataset in datasets:
        name = dataset.get('name', '')
        data = dataset.get('data', [])
        is_threshold = dataset.get('isThreshold', False)

        if is_threshold:
            ax.plot(x, data, '--', label=name, alpha=0.7, linewidth=1.5)
        else:
            ax.plot(x, data, '-', label=name, linewidth=2)

    if len(labels) > 10:
        step = len(labels) // 10
        ax.set_xticks(x[::step])
        ax.set_xticklabels([format_xaxis_label(l) for l in labels[::step]], rotation=45, ha='right', fontsize=8)
    else:
        ax.set_xticks(x)
        ax.set_xticklabels([format_xaxis_label(l) for l in labels], rotation=45, ha='right', fontsize=8)

    # 폰트 직접 적용
    if font_prop:
        ax.set_title(chart_title, fontsize=12, fontweight='bold', pad=10, fontproperties=font_prop)
    else:
        ax.set_title(chart_title, fontsize=12, fontweight='bold', pad=10)

    if locale == 'ko':
        xlabel = '시간'
        ylabel = '값'
    else:
        xlabel = 'Time'
        ylabel = 'Value'

    if font_prop:
        ax.set_xlabel(xlabel, fontsize=10, fontproperties=font_prop)
        ax.set_ylabel(ylabel, fontsize=10, fontproperties=font_prop)
    else:
        ax.set_xlabel(xlabel, fontsize=10)
        ax.set_ylabel(ylabel, fontsize=10)

    ax.grid(True, linestyle='--', alpha=0.3)
    ax.legend(loc='upper right', fontsize=8)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


# ============================================
# 상세 진단 섹션 생성
# ============================================
def create_diagnosis_detail_section(doc, main_data, detail_data, mode='diagnosis', locale='en'):
    """
    상세 진단 항목 섹션 생성
    """
    if mode == 'powerquality':
        if locale == 'ko':
            status_colors = {
                0: ('64748B', '정지', RGBColor(100, 116, 139)),
                1: ('10B981', '정상', RGBColor(16, 185, 129)),
                2: ('EAB308', '저', RGBColor(234, 179, 8)),
                3: ('F97316', '중', RGBColor(249, 115, 22)),
                4: ('EF4444', '고', RGBColor(239, 68, 68)),
            }
        else:
            status_colors = {
                0: ('64748B', 'Stop', RGBColor(100, 116, 139)),
                1: ('10B981', 'OK', RGBColor(16, 185, 129)),
                2: ('EAB308', 'Low', RGBColor(234, 179, 8)),
                3: ('F97316', 'Medium', RGBColor(249, 115, 22)),
                4: ('EF4444', 'High', RGBColor(239, 68, 68)),
            }
    else:
        if locale == 'ko':
            status_colors = {
                0: ('64748B', '정지', RGBColor(100, 116, 139)),
                1: ('10B981', '정상', RGBColor(16, 185, 129)),
                2: ('EAB308', '주의', RGBColor(234, 179, 8)),
                3: ('F97316', '경고', RGBColor(249, 115, 22)),
                4: ('EF4444', '위험', RGBColor(239, 68, 68)),
            }
        else:
            status_colors = {
                0: ('64748B', 'Stop', RGBColor(100, 116, 139)),
                1: ('10B981', 'OK', RGBColor(16, 185, 129)),
                2: ('EAB308', 'Warning', RGBColor(234, 179, 8)),
                3: ('F97316', 'Inspect', RGBColor(249, 115, 22)),
                4: ('EF4444', 'Repair', RGBColor(239, 68, 68)),
            }

    main_dict = {}
    for item in main_data:
        key = item.get('item_name', '').replace(' ', '')
        main_dict[key] = item

    grouped = {}
    for item in detail_data:
        parent_name = item.get('parent_name', '')
        if parent_name not in grouped:
            grouped[parent_name] = []
        grouped[parent_name].append(item)

    for parent_name, children in grouped.items():
        parent_key = parent_name.replace(' ', '')
        parent_info = main_dict.get(parent_key, {})
        parent_status = parent_info.get('status', 0)

        if parent_status < 2:
            continue

        if locale == 'ko':
            item_title = parent_info.get('title_ko') or parent_info.get('title') or parent_name
        elif locale == 'ja':
            item_title = parent_info.get('title_ja') or parent_info.get('title') or parent_name
        else:
            item_title = parent_info.get('title_en') or parent_info.get('title') or parent_name

        status_color, status_text, status_rgb = status_colors.get(int(parent_status), status_colors[0])

        header_table = doc.add_table(rows=1, cols=2)
        header_table.style = 'Table Grid'

        name_cell = header_table.rows[0].cells[0]
        name_cell.width = Inches(5)
        name_para = name_cell.paragraphs[0]
        name_run = name_para.add_run(item_title)
        name_run.font.size = Pt(12)
        name_run.font.bold = True

        status_cell = header_table.rows[0].cells[1]
        status_cell.width = Inches(1.5)
        status_para = status_cell.paragraphs[0]
        status_run = status_para.add_run(status_text)
        status_run.font.size = Pt(10)
        status_run.font.bold = True
        status_run.font.color.rgb = RGBColor(255, 255, 255)
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), status_color)
        status_cell._element.get_or_add_tcPr().append(shading_elm)

        problem_children = [c for c in children if c.get('status', 0) >= 2]

        if problem_children:
            detail_table = doc.add_table(rows=1, cols=3)
            detail_table.style = 'Light Grid Accent 1'

            if locale == 'ko':
                headers = ['항목', '모듈', '값']
            else:
                headers = ['Item', 'Module', 'Value']

            hdr_cells = detail_table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].paragraphs[0].add_run(header).bold = True

            for child in problem_children:
                if locale == 'ko':
                    child_title = child.get('title_ko') or child.get('title') or child.get('item_name', '')
                elif locale == 'ja':
                    child_title = child.get('title_ja') or child.get('title') or child.get('item_name', '')
                else:
                    child_title = child.get('title_en') or child.get('title') or child.get('item_name', '')

                row_cells = detail_table.add_row().cells
                row_cells[0].text = child_title
                row_cells[1].text = child.get('assembly_id', '')
                row_cells[2].text = str(child.get('value', ''))

        if locale == 'ko':
            description = parent_info.get('description_ko') or parent_info.get('description', '')
        elif locale == 'ja':
            description = parent_info.get('description_ja') or parent_info.get('description', '')
        else:
            description = parent_info.get('description_en') or parent_info.get('description', '')

        if description:
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(description)
            desc_run.font.size = Pt(10)
            desc_run.font.bold = True
            desc_run.font.color.rgb = status_rgb

        doc.add_paragraph()


# ============================================
# 트렌드 데이터 조회
# ============================================
async def get_trend_data_for_report(asset_name, detail_data, timestamp, mode='diagnosis'):
    """
    트렌드 차트용 데이터 조회
    """
    trend_list = []

    try:
        problem_items = []
        for item in detail_data:
            problem_items.append({
                'Name': item.get('item_name', ''),
                'AssemblyID': item.get('assembly_id', '')
            })

        if not problem_items:
            return trend_list

        param_type = 'powerquality' if mode == 'powerquality' else 'diagnostic'
        params_response = await get_params(asset_name, param_type, None)

        if not params_response.get('success'):
            logging.warning(f"파라미터 조회 실패: {params_response}")
            return trend_list

        param_data = params_response.get('data', [])

        matched_params = []
        matched_idx = set()

        for param in param_data:
            for item in problem_items:
                if param.get('Name') == item['Name'] and param.get('AssemblyID') == item['AssemblyID']:
                    param_id = param.get('ID')
                    if param_id and param_id not in matched_idx:
                        matched_params.append({
                            'idx': param_id,
                            'Assembly': param.get('AssemblyID', ''),
                            'title': param.get('Title', '')
                        })
                        matched_idx.add(param_id)

        if not matched_params:
            return trend_list

        from datetime import timedelta, timezone

        if 'Z' in timestamp:
            base_time = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
        else:
            base_time = datetime.fromisoformat(timestamp)

        if base_time.tzinfo is None:
            base_time = base_time.replace(tzinfo=timezone.utc)

        end_date = base_time
        start_date = base_time - timedelta(days=7)

        for param in matched_params:
            param_id = param['idx']
            title_name = f"[{param['Assembly']}]{param['title']}"

            trend_request = Trend(
                ParametersIds=[param_id],
                StartDate=start_date.strftime('%Y-%m-%dT%H:%M:%S'),
                EndDate=end_date.strftime('%Y-%m-%dT%H:%M:%S')
            )

            trend_response = await get_trendData(trend_request, None)

            if not trend_response.get('success'):
                continue

            res_data = trend_response.get('data', {})
            datasets = []
            labels = []

            for key, value in res_data.items():
                if key == 'Thresholds':
                    continue

                data_points = value.get('data', [])
                if data_points:
                    if not labels:
                        labels = [point.get('XAxis', '') for point in data_points]

                    datasets.append({
                        'name': value.get('Title', key),
                        'data': [point.get('YAxis', 0) for point in data_points],
                        'isThreshold': False
                    })

            thresholds = res_data.get('Thresholds', [])
            if thresholds and labels:
                threshold_strings = ["Out of Range(Down)", "Repair", "Inspect", "Warning", "Warning", "Inspect", "Repair", "Out of Range(Up)"]

                if thresholds and len(thresholds) > 0 and thresholds[0].get('Thresholds'):
                    threshold_count = len(thresholds[0]['Thresholds'])

                    for idx in range(threshold_count):
                        has_valid = any(
                            t.get('Thresholds', [None])[idx] not in [None, 'NaN']
                            and isinstance(t.get('Thresholds', [None])[idx], (int, float))
                            for t in thresholds
                        )

                        if not has_valid:
                            continue

                        threshold_data = []
                        for lbl in labels:
                            applicable_val = None

                            for t in sorted(thresholds, key=lambda x: x.get('XAxis', '')):
                                t_val = t.get('Thresholds', [None])[idx]
                                if t_val not in [None, 'NaN'] and isinstance(t_val, (int, float)):
                                    applicable_val = t_val

                            threshold_data.append(applicable_val if applicable_val is not None else 0)

                        if any(v for v in threshold_data):
                            datasets.append({
                                'name': threshold_strings[idx] if idx < len(threshold_strings) else f'Threshold {idx}',
                                'data': threshold_data,
                                'isThreshold': True
                            })

            if labels and datasets:
                trend_list.append({
                    'lineLabels': labels,
                    'lineData': datasets,
                    'lineTitle': title_name
                })

    except Exception as e:
        logging.error(f"트렌드 데이터 조회 실패: {e}")
        import traceback
        traceback.print_exc()

    return trend_list


# ============================================
# ITIC Curve 차트 생성 (기존 유지)
# ============================================
def create_itic_curve_chart(chart_path='/tmp/itic_curve.png', itic_events=None):
    """
    ITIC Curve 차트 생성
    """
    setup_korean_font()

    fig, ax = plt.subplots(figsize=(12, 8))

    prohibited_lower_x = [0.001, 0.01, 0.5, 10, 10, 0.001, 0.001]
    prohibited_lower_y = [0, 0, 70, 70, 0, 0, 0]
    ax.fill(prohibited_lower_x, prohibited_lower_y,
            color='#fee2e2', alpha=0.7, label='Prohibited Region')

    no_damage_x = [0.001, 0.01, 0.5, 10, 10, 0.5, 0.01, 0.001, 0.001]
    no_damage_y = [70, 70, 80, 80, 110, 120, 120, 110, 70]
    ax.fill(no_damage_x, no_damage_y,
            color='#fef3c7', alpha=0.7, label='No Damage')

    no_interrupt_x = [0.001, 0.01, 0.5, 10, 10, 0.5, 0.01, 0.001, 0.001]
    no_interrupt_y = [110, 120, 120, 110, 500, 500, 140, 140, 110]
    ax.fill(no_interrupt_x, no_interrupt_y,
            color='#d1fae5', alpha=0.7, label='No Interruption')

    prohibited_upper_x = [0.001, 0.01, 0.5, 10, 10, 0.001, 0.001]
    prohibited_upper_y = [140, 140, 500, 500, 200, 200, 140]
    ax.fill(prohibited_upper_x, prohibited_upper_y,
            color='#fee2e2', alpha=0.7)

    ax.plot([0.001, 0.01, 0.5, 10], [70, 70, 80, 80],
            'r-', linewidth=2, label='ITIC Limit')
    ax.plot([0.001, 0.01, 0.5, 10], [110, 120, 120, 110],
            'r-', linewidth=2)
    ax.plot([0.001, 0.01, 0.5], [140, 140, 500],
            'r-', linewidth=2)

    if itic_events:
        for event in itic_events:
            duration = event.get('duration', 0)
            voltage_pct = event.get('voltage_pct', 100)
            event_type = event.get('event_type', 'UNKNOWN')

            color = 'orange' if event_type == 'SAG' else 'green'

            ax.scatter(duration, voltage_pct,
                       c=color, s=50, alpha=0.8,
                       edgecolors='black', linewidth=1,
                       zorder=5)

    ax.set_xscale('log')
    ax.set_xlim(0.001, 10)
    ax.set_ylim(0, 200)

    ax.set_xlabel('Duration (s)', fontsize=12, fontweight='bold')
    ax.set_ylabel('Voltage (%)', fontsize=12, fontweight='bold')
    ax.set_title('ITIC Curve - Power Quality Events', fontsize=14, fontweight='bold', pad=20)

    ax.grid(True, which='both', alpha=0.3, linestyle='--')
    ax.legend(loc='upper right', fontsize=10)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()

    print(f"✅ ITIC Curve 차트 저장: {chart_path}")
    return chart_path


# ============================================
# 기존 리포트 생성 (EN50160)
# ============================================
async def generate_report(asset_name, asset_type, location='', output_path='/home/root/logs/report.docx', itic_events=None):
    print(f"📡 API 호출: {asset_name}")
    api_response = await get_asset(asset_name)

    if not api_response.get('success'):
        print(f"❌ 설비 정보 가져오기 실패: {api_response.get('error')}")
        return None

    parsed_data = parse_asset_data(api_response)

    channel_data = {
        'channel': 'Main',
        'name': asset_name,
        'type': asset_type,
        'drive_type': parsed_data.get('drive_type', ''),
        'location': location,
        'mac': parsed_data.get('mac', '-'),
        'rated_voltage': parsed_data.get('rated_voltage', '-'),
        'rated_current': parsed_data.get('rated_current', '-'),
        'rated_power': parsed_data.get('rated_power', '-'),
        'rated_capacity': parsed_data.get('rated_capacity', '-'),
        'frequency': parsed_data.get('frequency', 60),
        'pt_wiring_mode': parsed_data.get('pt_wiring_mode', '-'),
    }

    print(f"✅ 설비 정보 파싱 완료")

    doc = Document()

    style = doc.styles['Normal']
    style.font.size = Pt(10)

    title = doc.add_heading('Power Quality Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('EN50160 Weekly Report')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'\nGenerated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    date_run.font.size = Pt(11)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading('1. Asset Information', 1)
    create_channel_info_section(doc, channel_data)

    doc.add_page_break()

    doc.add_heading('2. ITIC Curve Analysis', 1)

    intro = doc.add_paragraph(
        'The ITIC (Information Technology Industry Council) Curve defines the voltage '
        'tolerance limits for IT equipment. This chart shows power quality events '
        'recorded during the measurement period.'
    )
    intro.paragraph_format.space_after = Pt(12)

    chart_path = create_itic_curve_chart(itic_events=itic_events)
    doc.add_picture(chart_path, width=Inches(6.5))

    if itic_events:
        doc.add_paragraph()
        analysis = doc.add_paragraph()
        analysis_title = analysis.add_run('Analysis Results:\n')
        analysis_title.font.bold = True
        analysis_title.font.size = Pt(11)

        sag_count = sum(1 for e in itic_events if e.get('event_type') == 'SAG')
        swell_count = sum(1 for e in itic_events if e.get('event_type') == 'SWELL')

        analysis_content = analysis.add_run(
            f'• Total Events: {len(itic_events)}\n'
            f'• Voltage Sag (SAG): {sag_count}\n'
            f'• Voltage Swell (SWELL): {swell_count}\n'
            f'• All events are within ITIC acceptable limits.'
        )

    doc.save(output_path)

    return output_path


# ============================================
# API 라우터
# ============================================
@router.post("/generate")
async def set_report(request: Request):
    data = await request.json()
    o_path = await generate_report(data["assetName"], data["assetType"], data["location"])
    print(f"\n✅ 리포트 생성 완료: {o_path}")
    return True


@router.get('/lastReportData/{mode}/{asset_name}')
async def get_last_diagnosis(mode: str, asset_name: str):
    """마지막 진단 데이터 조회"""
    try:
        query_api = influx_state.query_api

        logging.info(f"🔍 조회 시작: mode={mode}, asset_name={asset_name}")

        timestamp_query = f'''
        from(bucket: "ntek")
            |> range(start: -7d)
            |> filter(fn: (r) => r["_measurement"] == "{mode}")
            |> filter(fn: (r) => r["asset_name"] == "{asset_name}")
            |> filter(fn: (r) => r["data_type"] == "main")
            |> filter(fn: (r) => r["_field"] == "status")
            |> last()
        '''

        ts_tables = query_api.query(timestamp_query)

        last_time = None
        for table in ts_tables:
            for record in table.records:
                t = record.get_time()
                if last_time is None or t > last_time:
                    last_time = t

        if not last_time:
            logging.warning("⚠️ 데이터 없음")
            return {"success": False, "msg": "No Data"}

        logging.info(f"📅 마지막 타임스탬프: {last_time.isoformat()}")

        from datetime import timedelta, timezone
        start_time = (last_time - timedelta(seconds=1)).strftime('%Y-%m-%dT%H:%M:%S.%fZ')
        end_time = (last_time + timedelta(seconds=1)).strftime('%Y-%m-%dT%H:%M:%S.%fZ')

        data_query = f'''
        from(bucket: "ntek")
            |> range(start: {start_time}, stop: {end_time})
            |> filter(fn: (r) => r["_measurement"] == "{mode}")
            |> filter(fn: (r) => r["asset_name"] == "{asset_name}")
        '''

        logging.info(f"📊 데이터 쿼리 실행")
        tables = query_api.query(data_query)

        grouped = {"main": {}, "detail": {}}

        for table in tables:
            for record in table.records:
                data_type = record.values.get("data_type")
                item_name = record.values.get("item_name")

                if not data_type or data_type not in ["main", "detail"]:
                    continue
                if not item_name:
                    continue

                if item_name not in grouped[data_type]:
                    grouped[data_type][item_name] = {
                        "item_name": item_name,
                        "data_type": data_type,
                        "asset_name": record.values.get("asset_name"),
                        "channel": record.values.get("channel"),
                    }
                    if data_type == "detail":
                        grouped[data_type][item_name]["parent_name"] = record.values.get("parent_name")

                field_name = record.get_field()
                field_value = record.get_value()
                grouped[data_type][item_name][field_name] = field_value

        main_data = list(grouped["main"].values())
        detail_data = list(grouped["detail"].values())

        if last_time.tzinfo is not None:
            last_time_local = last_time.astimezone().replace(tzinfo=None)
        else:
            last_time_local = last_time

        logging.info(f"✅ 최종 결과: main={len(main_data)}개, detail={len(detail_data)}개")

        result = {
            "asset_name": asset_name,
            "mode": mode,
            "timestamp": last_time_local.isoformat() if last_time_local else None,
            "main": main_data,
            "detail": detail_data
        }

        return {"success": True, "data": result}

    except Exception as e:
        logging.error(f"❌ 마지막 진단 데이터 조회 실패: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "msg": str(e)}


@router.get('/reportTimes/{date}/{asset_name}/{mode}')
async def get_report_times(date: str, asset_name: str, mode: str):
    """특정 날짜의 저장 시간 목록 조회 (mode별)"""
    try:
        query_api = influx_state.query_api

        start_time = f"{date}T00:00:00Z"
        end_time = f"{date}T23:59:59Z"

        query = f'''
        from(bucket: "ntek")
            |> range(start: {start_time}, stop: {end_time})
            |> filter(fn: (r) => r["_measurement"] == "{mode}")
            |> filter(fn: (r) => r["asset_name"] == "{asset_name}")
            |> filter(fn: (r) => r["data_type"] == "main")
            |> filter(fn: (r) => r["_field"] == "status")
        '''

        tables = query_api.query(query)

        times = set()
        for table in tables:
            for record in table.records:
                t = record.get_time()
                if t.tzinfo is not None:
                    t_local = t.astimezone().replace(tzinfo=None)
                else:
                    t_local = t
                times.add(t_local.isoformat())

        sorted_times = sorted(list(times), reverse=True)

        logging.info(f"📅 {mode} {date} 시간 목록: {len(sorted_times)}개")
        return {"success": True, "data": sorted_times}

    except Exception as e:
        logging.error(f"❌ 시간 목록 조회 실패: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "msg": str(e)}


@router.get('/reportDataByTime/{mode}/{asset_name}/{timestamp}')
async def get_report_data_by_time(mode: str, asset_name: str, timestamp: str):
    """특정 시간의 데이터 조회"""
    try:
        query_api = influx_state.query_api

        from datetime import timedelta, timezone
        import time as time_module

        local_time = datetime.fromisoformat(timestamp)
        if local_time.tzinfo is None:
            utc_offset = -time_module.timezone
            local_tz = timezone(timedelta(seconds=utc_offset))
            local_time = local_time.replace(tzinfo=local_tz)

        utc_time = local_time.astimezone(timezone.utc)

        start_time = (utc_time - timedelta(seconds=1)).strftime('%Y-%m-%dT%H:%M:%S.%fZ')
        end_time = (utc_time + timedelta(seconds=1)).strftime('%Y-%m-%dT%H:%M:%S.%fZ')

        logging.info(f"🔍 조회: mode={mode}, asset={asset_name}, time={timestamp}")

        data_query = f'''
        from(bucket: "ntek")
            |> range(start: {start_time}, stop: {end_time})
            |> filter(fn: (r) => r["_measurement"] == "{mode}")
            |> filter(fn: (r) => r["asset_name"] == "{asset_name}")
        '''

        tables = query_api.query(data_query)

        grouped = {"main": {}, "detail": {}}

        for table in tables:
            for record in table.records:
                data_type = record.values.get("data_type")
                item_name = record.values.get("item_name")

                if not data_type or data_type not in ["main", "detail"]:
                    continue
                if not item_name:
                    continue

                if item_name not in grouped[data_type]:
                    grouped[data_type][item_name] = {
                        "item_name": item_name,
                        "data_type": data_type,
                        "asset_name": record.values.get("asset_name"),
                        "channel": record.values.get("channel"),
                    }
                    if data_type == "detail":
                        grouped[data_type][item_name]["parent_name"] = record.values.get("parent_name")
                        grouped[data_type][item_name]["assembly_id"] = record.values.get("assembly_id")

                field_name = record.get_field()
                field_value = record.get_value()
                grouped[data_type][item_name][field_name] = field_value

        main_data = list(grouped["main"].values())
        detail_data = list(grouped["detail"].values())

        logging.info(f"✅ 결과: main={len(main_data)}개, detail={len(detail_data)}개")

        result = {
            "asset_name": asset_name,
            "mode": mode,
            "timestamp": timestamp,
            "main": main_data,
            "detail": detail_data
        }

        return {"success": True, "data": result}

    except Exception as e:
        logging.error(f"❌ 데이터 조회 실패: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "msg": str(e)}


@router.get('/status_trend/{mode}/{asset_name}/{item_name}')
async def get_item_trend(mode: str, asset_name: str, item_name: str):
    """특정 항목의 2주치 트렌드 조회"""
    try:
        query_api = influx_state.query_api

        logging.info(f"🔍 트렌드 조회: mode={mode}, asset_name={asset_name}, item={item_name}")

        query = f'''
        from(bucket: "ntek")
            |> range(start: -14d)
            |> filter(fn: (r) => r["_measurement"] == "{mode}")
            |> filter(fn: (r) => r["asset_name"] == "{asset_name}")
            |> filter(fn: (r) => r["data_type"] == "main")
            |> filter(fn: (r) => r["item_name"] == "{item_name}")
            |> filter(fn: (r) => r["_field"] == "status")
            |> sort(columns: ["_time"])
        '''

        tables = query_api.query(query)

        trend = []
        for table in tables:
            for record in table.records:
                t = record.get_time()
                if t.tzinfo is not None:
                    t = t.astimezone().replace(tzinfo=None)

                trend.append({
                    "timestamp": t.isoformat(),
                    "status": record.get_value()
                })

        logging.info(f"✅ 트렌드 결과: {len(trend)}개")

        return {
            "success": True,
            "data": {
                "asset_name": asset_name,
                "mode": mode,
                "item_name": item_name,
                "count": len(trend),
                "trend": trend
            }
        }

    except Exception as e:
        logging.error(f"❌ 트렌드 조회 실패: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "msg": str(e)}

# ============================================
# 진단 리포트 워드 다운로드 API
# ============================================
@router.get("/downloadDiagnosisReport/{mode}/{asset_name}/{channel}/{timestamp:path}")
async def download_diagnosis_report(mode: str, asset_name: str, channel: str, timestamp: str, locale: str = 'en'):
    """
    진단 리포트 워드 파일 다운로드 (GET)
    """
    try:
        logging.info(f"📄 리포트 생성 요청: {mode}, {asset_name}, {channel}, {timestamp}, locale={locale}")

        report_response = await get_report_data_by_time(mode, asset_name, timestamp)

        if not report_response.get('success'):
            raise HTTPException(status_code=404, detail="Report data not found")

        report_data = report_response.get('data', {})
        main_data = report_data.get('main', [])
        detail_data = report_data.get('detail', [])

        api_response = await get_asset(asset_name)
        parsed_data = parse_asset_data(api_response) if api_response.get('success') else {}

        trend_data = await get_trend_data_for_report(asset_name, detail_data, timestamp, mode)
        logging.info(f"📈 트렌드 차트 {len(trend_data)}개 조회됨")

        temp_dir = tempfile.gettempdir()
        report_id = str(uuid.uuid4())[:8]
        output_path = f'{temp_dir}/diagnosis_report_{report_id}.docx'

        doc = Document()

        # 문서 여백 설정 (좁게)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        style = doc.styles['Normal']
        style.font.size = Pt(10)

        # === 1. 표지 ===
        if locale == 'ko':
            title_text = '전력품질 진단 리포트' if mode == 'powerquality' else '설비 진단 리포트'
        elif locale == 'ja':
            title_text = '電力品質診断レポート' if mode == 'powerquality' else '設備診断レポート'
        else:
            title_text = 'Power Quality Diagnosis Report' if mode == 'powerquality' else 'Equipment Diagnosis Report'
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        asset_para = doc.add_paragraph()
        asset_run = asset_para.add_run(f'{asset_name} ({channel})')
        asset_run.font.size = Pt(18)
        asset_run.font.bold = True
        asset_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph()
        if locale == 'ko':
            date_label = '리포트 날짜'
        else:
            date_label = 'Report Date'
        date_run = date_para.add_run(f'\n{date_label}: {timestamp}')
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        generated_para = doc.add_paragraph()
        if locale == 'ko':
            gen_label = '생성일시'
        else:
            gen_label = 'Generated'
        generated_run = generated_para.add_run(f'{gen_label}: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        generated_run.font.size = Pt(10)
        generated_run.font.color.rgb = RGBColor(100, 116, 139)
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # === 2. 설비 정보 ===
        if locale == 'ko':
            section1_title = '1. 설비 정보'
        else:
            section1_title = '1. Asset Information'
        doc.add_heading(section1_title, 1)

        if api_response.get('success'):
            asset_type = api_response.get('driveType', 'Motor')
            channel_data = {
                'channel': channel,
                'name': asset_name,
                'type': asset_type,
                'drive_type': parsed_data.get('drive_type', ''),
                'location': '',
                'mac': parsed_data.get('mac', '-'),
                'rated_voltage': parsed_data.get('rated_voltage', '-'),
                'rated_current': parsed_data.get('rated_current', '-'),
                'rated_power': parsed_data.get('rated_power', '-'),
                'rated_capacity': parsed_data.get('rated_capacity', '-'),
                'frequency': parsed_data.get('frequency', 60),
                'pt_wiring_mode': parsed_data.get('pt_wiring_mode', '-'),
            }
            create_channel_info_section(doc, channel_data, locale)

        doc.add_page_break()

        # === 3. 진단 결과 바차트 ===
        if mode == 'powerquality':
            if locale == 'ko':
                section2_title = '2. 전력품질 진단 현황'
            else:
                section2_title = '2. Power Quality Status'
        else:
            if locale == 'ko':
                section2_title = '2. 설비 진단 현황'
            else:
                section2_title = '2. Diagnostic Status'
        doc.add_heading(section2_title, 1)

        chart_path = f'{temp_dir}/diagnosis_chart_{report_id}.png'
        create_diagnosis_bar_chart(chart_path, main_data, mode, locale)

        if os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(7.3))
            os.remove(chart_path)

        doc.add_page_break()

        # === 4. 상세 항목 설명 ===
        if locale == 'ko':
            section3_title = '3. 상세 분석'
        else:
            section3_title = '3. Detail Analysis'
        doc.add_heading(section3_title, 1)

        has_issues = any(item.get('status', 0) >= 2 for item in main_data)

        if has_issues:
            create_diagnosis_detail_section(doc, main_data, detail_data, mode, locale)
        else:
            if locale == 'ko':
                doc.add_paragraph('문제가 발견되지 않았습니다. 모든 항목이 정상 범위 내에 있습니다.')
            else:
                doc.add_paragraph('No issues detected. All items are within normal range.')

        # === 5. 트렌드 차트 ===
        if trend_data:
            doc.add_page_break()
            if locale == 'ko':
                section4_title = '4. 트렌드 차트'
            else:
                section4_title = '4. Trend Charts'
            doc.add_heading(section4_title, 1)

            for idx, trend in enumerate(trend_data):
                if not trend.get('lineLabels') or not trend.get('lineData'):
                    continue

                trend_chart_path = f'{temp_dir}/trend_chart_{report_id}_{idx}.png'
                result = create_trend_chart(trend_chart_path, trend, locale)

                if result and os.path.exists(trend_chart_path):
                    doc.add_picture(trend_chart_path, width=Inches(7.3))
                    doc.add_paragraph()
                    os.remove(trend_chart_path)

        doc.save(output_path)
        logging.info(f"✅ 진단 리포트 생성 완료: {output_path}")

        date_str = timestamp.split('T')[0] if 'T' in timestamp else timestamp[:10]
        filename = f'{mode}_report_{asset_name}_{date_str}.docx'

        return FileResponse(
            path=output_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"❌ 리포트 생성 실패: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


async def get_en50160_report_by_date(channel: str, date: str):
    try:
        from datetime import datetime, timedelta

        # 날짜 파싱
        target_date = datetime.strptime(date, "%Y%m%d")
        start_of_day = target_date.strftime("%Y-%m-%dT00:00:00Z")
        end_of_day = (target_date + timedelta(days=1)).strftime("%Y-%m-%dT00:00:00Z")

        query = f'''
        from(bucket: "ntek30")
            |> range(start: {start_of_day}, stop: {end_of_day})
            |> filter(fn: (r) => r["_measurement"] == "en50160")
            |> filter(fn: (r) => r["channel"] == "{channel}")
            |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
        '''

        tables = influx_state.query_api.query(query)

        for table in tables:
            for record in table.records:
                return {"success": True, "data": dict(record.values)}

        return {"success": False, "msg": f"리포트 없음: {channel}/{date}"}

    except Exception as e:
        logging.error(f"EN50160 리포트 조회 실패: {e}")
        return {"success": False, "msg": str(e)}

# @router.get("/getEn50160_summary/{channel}/{date}")
# async def get_summaryData(channel, date):
#     try:
#         result = await get_en50160_report_by_date(channel, date)
#
#         if result is None:
#             return {"success": False, "msg": f"리포트 없음: {channel}/{date}"}
#
#         return {"success": True, "data": result}
#
#     except Exception as e:
#         logging.error(f"EN50160 리포트 조회 실패: {e}")
#         return {"success": False, "msg": str(e)}

@router.get("/week/{channel}/{filename}")
async def get_weekly_report(channel:str, filename: str):
    # get_all_chart_data 사용 (내부에서 파일 1번만 읽음)
    set_limit(channel)
    return processor.get_all_chart_data(channel,filename=filename)


def set_limit(channel):
    if redis_state.client.hexists("Equipment","ChannelData"):
        chDict = json.loads(redis_state.client.hget("Equipment","ChannelData"))
        if channel == 'Main' or channel == 'main':
            chName = 'main'
        else:
            chName = 'sub'
        processor.set_limits(float(chDict[chName]["RatedVoltage"]), float(chDict[chName]["RatedCurrent"]),float(chDict[chName]["RatedFrequency"]))
    else:
        setting = json.loads(redis_state.client.hget("System","setup"))
        for ch in setting.get("channel", []):
            name = ch.get("channel", "")
            if name == "Main":
                processor.set_limits(float(ch[name]["ptInfo"]["vnorminal"]), float(ch[channel]["ctInfo"]["inorminal"]),float(ch[channel]["ptInfo"]["linefrequency"]))
            elif name == "Sub":
                processor.set_limits(float(ch[name]["ptInfo"]["vnorminal"]), float(ch[channel]["ctInfo"]["inorminal"]),float(ch[channel]["ptInfo"]["linefrequency"]))

@router.get("/list/{channel}")
def get_filelist(channel):
    filelist = processor.list_files(channel)
    return {"success": True, "data" : filelist}


@router.get('/getReportDiagnosis/{mode}/{asset}/{channel}/{date}')
async def get_report_diagnosis_data(mode: str, asset: str, channel: str, date: str):
    """
    저장된 Parquet 파일에서 진단 데이터 + 트렌드 한꺼번에 조회
    Args:
        mode: diagnosis 또는 powerquality
        channel: Main 또는 Sub
        date: 날짜 (20251217 형식)
    """
    try:
        import pandas as pd
        import json
        import numpy as np
        from pathlib import Path

        def convert_to_serializable(obj):
            """numpy/pandas 타입을 JSON 직렬화 가능한 타입으로 변환"""
            if isinstance(obj, dict):
                return {k: convert_to_serializable(v) for k, v in obj.items()}
            elif isinstance(obj, list):
                return [convert_to_serializable(v) for v in obj]
            elif isinstance(obj, (np.integer,)):
                return int(obj)
            elif isinstance(obj, (np.floating,)):
                return float(obj)
            elif isinstance(obj, np.ndarray):
                return obj.tolist()
            elif isinstance(obj, (pd.Timestamp,)):
                return obj.isoformat()
            elif pd.isna(obj):
                return None
            return obj

        # 파일 경로 구성
        output_dir = Path("/usr/local/sv500/reports") / channel;
        pattern = f"diagnosis_report_{asset}_{date}.parquet"
        files = list(output_dir.glob(pattern))

        if not files:
            return {"success": False, "msg": f"파일 없음: {pattern}"}

        filepath = files[0]
        logging.info(f"📂 파일 로드: {filepath.name}")

        # Parquet 읽기
        df = pd.read_parquet(filepath)
        record = df.iloc[0].to_dict()

        # mode에 해당하는 데이터 추출
        mode_data = record.get(mode)

        if mode_data is None:
            return {"success": False, "msg": f"{mode} 데이터 없음"}

        # JSON 문자열이면 파싱
        if isinstance(mode_data, str):
            mode_data = json.loads(mode_data)

        main_data = mode_data.get("main", []) if mode_data else []
        detail_data = mode_data.get("detail", []) if mode_data else []
        trends_data = mode_data.get("trends") if mode_data else None

        # 각각 문자열이면 파싱
        if isinstance(main_data, str):
            main_data = json.loads(main_data)
        if isinstance(detail_data, str):
            detail_data = json.loads(detail_data)
        if isinstance(trends_data, str):
            trends_data = json.loads(trends_data)

        result = {
            "asset_name": record.get("asset_name"),
            "mode": mode,
            "timestamp": mode_data.get("timestamp") if mode_data else None,
            "main": main_data,
            "detail": detail_data,
            "trends": trends_data
        }

        # numpy/pandas 타입 변환
        result = convert_to_serializable(result)

        trends_count = len(result['trends']) if result['trends'] else 0
        logging.info(f"✅ 결과: main={len(result['main'])}개, detail={len(result['detail'])}개, trends={trends_count}개 항목")

        return {"success": True, "data": result}

    except Exception as e:
        logging.error(f"❌ 데이터 조회 실패: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "msg": str(e)}


# ============================================
# 유틸리티 함수
# ============================================
def convert_to_serializable(obj):
    """numpy/pandas 타입을 JSON 직렬화 가능한 타입으로 변환"""
    if isinstance(obj, dict):
        return {k: convert_to_serializable(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_to_serializable(v) for v in obj]
    elif isinstance(obj, (np.integer,)):
        return int(obj)
    elif isinstance(obj, (np.floating,)):
        return float(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, (pd.Timestamp,)):
        return obj.isoformat()
    elif pd.isna(obj):
        return None
    return obj


def parse_json_field(value):
    """JSON 문자열이면 파싱, 아니면 그대로 반환"""
    if isinstance(value, str):
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            return value
    return value


def get_week_range_from_date(date_str: str):
    """
    날짜 문자열(20251217)로부터 해당 주의 시작/종료일 계산

    Returns:
        (start_date, end_date) - 둘 다 datetime 객체
    """
    # 20251217 → 2025-12-17
    target_date = datetime.strptime(date_str, "%Y%m%d")

    # 해당 주의 월요일 찾기
    monday = target_date - timedelta(days=target_date.weekday())
    sunday = monday + timedelta(days=6)

    return monday, sunday


# ============================================
# 데이터 로드 함수
# ============================================
def get_en50160_from_redis(channel: str, existing_data: Optional[Dict] = None) -> Optional[Dict]:
    """
    EN50160 요약 테이블 데이터 조회 (API 호출)

    Args:
        channel: 채널명 (Main/Sub)
        existing_data: 기존 EN50160 데이터 (있으면 detail_table만 업데이트)

    Returns:
        EN50160 데이터 딕셔너리
    """
    try:
        from routes.api import getEn50160

        # API 호출로 데이터 가져오기
        result = getEn50160(channel)

        if result and result.get('success'):
            tbdata = result.get('data', {})

            if tbdata:
                # 테이블 데이터 변환
                detail_table = transform_en50160_table_data(tbdata)

                if existing_data:
                    existing_data['detail_table'] = detail_table
                    return existing_data
                else:
                    return {'detail_table': detail_table}

    except ImportError:
        logging.warning("getEn50160 import 실패")
    except Exception as e:
        logging.error(f"EN50160 데이터 조회 실패: {e}")

    return existing_data


def transform_en50160_table_data(tbdata: Dict) -> List[Dict]:
    """
    getEn50160 API 데이터를 테이블 형식으로 변환
    """
    # 파라미터 정의
    params_def = [
        {'parameter': 'Frequency Variation 1', 'parameter_ko': '주파수 변동 1',
         'keys': {'l1': 'Frequency Variation 1', 'l2': None, 'l3': None},
         'required': '(+1%~-1%), 99.5%/Week', 'bits': [0]},
        {'parameter': 'Frequency Variation 2', 'parameter_ko': '주파수 변동 2',
         'keys': {'l1': 'Frequency Variation 2', 'l2': None, 'l3': None},
         'required': '(+4%~-6%), 100%/Week', 'bits': [1]},
        {'parameter': 'Voltage Variation 1', 'parameter_ko': '전압 변동 1',
         'keys': {'l1': 'Voltage Variation 1 L1(%)', 'l2': 'Voltage Variation 1 L2(%)',
                  'l3': 'Voltage Variation 1 L3(%)'},
         'required': '(+10%~-10%), 95%/Week', 'bits': [2, 3, 4]},
        {'parameter': 'Voltage Variation 2', 'parameter_ko': '전압 변동 2',
         'keys': {'l1': 'Voltage Variation 2 L1(%)', 'l2': 'Voltage Variation 2 L2(%)',
                  'l3': 'Voltage Variation 2 L3(%)'},
         'required': '(+10%~-15%), 100%/Week', 'bits': [5, 6, 7]},
        {'parameter': 'Voltage Unbalance', 'parameter_ko': '전압 불평형',
         'keys': {'l1': 'Voltage Unbalance', 'l2': None, 'l3': None},
         'required': '(2% 미만), 100%/Week', 'bits': [8]},
        {'parameter': 'THD', 'parameter_ko': 'THD',
         'keys': {'l1': 'THDs Variation L1(%)', 'l2': 'THDs Variation L2(%)', 'l3': 'THDs Variation L3(%)'},
         'required': '(8%미만), 95%/Week', 'bits': [9, 10, 11]},
        {'parameter': 'Harmonics', 'parameter_ko': '고조파',
         'keys': {'l1': 'Harmonics Variatiopn L1(%)', 'l2': 'Harmonics Variatiopn L2(%)',
                  'l3': 'Harmonics Variatiopn L3(%)'},
         'required': '(0.5%~6%), 95%/Week', 'bits': [12, 13, 14]},
        {'parameter': 'Pst', 'parameter_ko': 'Pst',
         'keys': {'l1': 'Flickers Pst L1(%)', 'l2': 'Flickers Pst L2(%)', 'l3': 'Flickers Pst L3(%)'},
         'required': '(1), Info Only', 'bits': [15, 16, 17]},
        {'parameter': 'Plt', 'parameter_ko': 'Plt',
         'keys': {'l1': 'Flickers Plt L1(%)', 'l2': 'Flickers Plt L2(%)', 'l3': 'Flickers Plt L3(%)'},
         'required': '(1), 95%/Week', 'bits': [18, 19, 20]},
        {'parameter': 'Signal Vol.', 'parameter_ko': '신호 전압',
         'keys': {'l1': 'Signaling Voltage L1(%)', 'l2': 'Signaling Voltage L2(%)', 'l3': 'Signaling Voltage L3(%)'},
         'required': '(10%미만), 99%/Day', 'bits': [21, 22, 23]},
        {'parameter': 'Voltage Sag', 'parameter_ko': '전압 강하',
         'keys': {'l1': 'Voltage Dips L1', 'l2': 'Voltage Dips L2', 'l3': 'Voltage Dips L3'},
         'required': 'Info Only', 'bits': None},
        {'parameter': 'Voltage Swell', 'parameter_ko': '전압 상승',
         'keys': {'l1': 'Voltage Swells L1', 'l2': 'Voltage Swells L2', 'l3': 'Voltage Swells L3'},
         'required': 'Info Only', 'bits': None},
        {'parameter': 'Short Interruption', 'parameter_ko': '단시간 정전',
         'keys': {'l1': 'Short Interruptions L1', 'l2': 'Short Interruptions L2', 'l3': 'Short Interruptions L3'},
         'required': 'Info Only', 'bits': None},
        {'parameter': 'Long Interruption', 'parameter_ko': '장시간 정전',
         'keys': {'l1': 'Long Interruptions L1', 'l2': 'Long Interruptions L2', 'l3': 'Long Interruptions L3'},
         'required': 'Info Only', 'bits': None},
        {'parameter': 'Signaling Volt.', 'parameter_ko': '신호 전압',
         'keys': {'l1': 'Signaling Voltage L1(%)', 'l2': 'Signaling Voltage L2(%)', 'l3': 'Signaling Voltage L3(%)'},
         'required': 'Info Only', 'bits': None},
    ]

    # status&compliance 비트마스크
    compliance_mask = tbdata.get('status&compliance', 0)

    detail_table = []
    for param in params_def:
        keys = param['keys']

        # L1, L2, L3 값 가져오기
        l1 = tbdata.get(keys['l1'], '-') if keys['l1'] else '-'
        l2 = tbdata.get(keys['l2'], '-') if keys['l2'] else '-'
        l3 = tbdata.get(keys['l3'], '-') if keys['l3'] else '-'

        # Compliance 판정
        if param['bits']:
            failed = any((compliance_mask & (1 << bit)) != 0 for bit in param['bits'])
            compliance = 'Failed' if failed else 'OK'
        else:
            compliance = '-'

        detail_table.append({
            'parameter': param['parameter'],
            'parameter_ko': param['parameter_ko'],
            'l1': l1,
            'l2': l2,
            'l3': l3,
            'compliance': compliance,
            'required': param['required']
        })

    return detail_table


def get_itic_data_for_report(channel: str) -> Optional[List[Dict]]:
    """
    ITIC 이벤트 데이터 조회

    Args:
        channel: 채널명 (Main/Sub)

    Returns:
        ITIC 이벤트 리스트 [{duration, voltage_pct, event_type}, ...]
    """
    try:
        from routes.api import getITIC_data, format_influx_time, get_setupContext

        data = getITIC_data("events", channel)

        if not data or len(data) == 0:
            return None

        # ratedV 조회
        ch = 'main' if channel == 'Main' else 'sub'
        setupdict = get_setupContext()
        if not setupdict:
            return None

        rated_itic = setupdict[ch]
        rated_voltage = rated_itic["ptInfo"]["vnorminal"]

        # ITIC 이벤트 형식으로 변환
        itic_events = []

        def parse_mask(mask):
            """마스크에서 위상 리스트 파싱"""
            phases = []
            if mask & 1: phases.append(0)  # L1
            if mask & 2: phases.append(1)  # L2
            if mask & 4: phases.append(2)  # L3
            return phases if phases else [0]

        def get_fin_value(phaselist, levellist, event_type):
            """위상별 레벨에서 최종 값 계산"""
            values = [levellist[p] for p in phaselist if p < len(levellist) and levellist[p] is not None]
            if not values:
                return 0
            # SAG는 최소값, SWELL은 최대값
            return min(values) if event_type == 'SAG' else max(values)

        for event in data:
            phaselist = parse_mask(event.get("mask", 7))
            levellist = [event.get("level_l1"), event.get("level_l2"), event.get("level_l3")]

            voltage_value = get_fin_value(phaselist, levellist, event.get("event_type", "SAG"))
            voltage_pct = (voltage_value * 100) / rated_voltage if rated_voltage else 0
            duration_sec = event.get("duration", 0) / 1000.0  # ms -> sec

            itic_events.append({
                "duration": duration_sec,
                "voltage_pct": round(voltage_pct, 2),
                "event_type": event.get("event_type", "SAG"),
                "event_time": format_influx_time(event.get("time")) if event.get("time") else None
            })

        logging.info(f"ITIC 이벤트 {len(itic_events)}개 조회됨")
        return itic_events

    except ImportError as e:
        logging.warning(f"ITIC 관련 모듈 import 실패: {e}")
        return None
    except Exception as e:
        logging.error(f"ITIC 데이터 조회 실패: {e}")
        return None


def transform_en50160_data(harmdict: Dict) -> Dict:
    summary = {
        'compliance': 'PASS',  # 기본값
        'frequency': {'result': 'N/A'},
        'voltage': {'result': 'N/A'},
        'thd': {'result': 'N/A'},
        'unbalance': {'result': 'N/A'},
        'flicker': {'result': 'N/A'},
        'harmonics': {'result': 'N/A'}
    }

    try:
        # harmdict 구조에 따라 매핑
        if 'frequency' in harmdict:
            freq_data = harmdict['frequency']
            if isinstance(freq_data, dict):
                summary['frequency']['result'] = freq_data.get('result', freq_data.get('status', 'N/A'))
            else:
                summary['frequency']['result'] = 'PASS' if freq_data else 'FAIL'

        if 'voltage' in harmdict:
            volt_data = harmdict['voltage']
            if isinstance(volt_data, dict):
                summary['voltage']['result'] = volt_data.get('result', volt_data.get('status', 'N/A'))
            else:
                summary['voltage']['result'] = 'PASS' if volt_data else 'FAIL'

        if 'thd' in harmdict:
            thd_data = harmdict['thd']
            if isinstance(thd_data, dict):
                summary['thd']['result'] = thd_data.get('result', thd_data.get('status', 'N/A'))
            else:
                summary['thd']['result'] = 'PASS' if thd_data else 'FAIL'

        if 'unbalance' in harmdict:
            unbal_data = harmdict['unbalance']
            if isinstance(unbal_data, dict):
                summary['unbalance']['result'] = unbal_data.get('result', unbal_data.get('status', 'N/A'))
            else:
                summary['unbalance']['result'] = 'PASS' if unbal_data else 'FAIL'

        if 'flicker' in harmdict:
            flicker_data = harmdict['flicker']
            if isinstance(flicker_data, dict):
                summary['flicker']['result'] = flicker_data.get('result', flicker_data.get('status', 'N/A'))
            else:
                summary['flicker']['result'] = 'PASS' if flicker_data else 'FAIL'

        if 'harmonics' in harmdict:
            harm_data = harmdict['harmonics']
            if isinstance(harm_data, dict):
                summary['harmonics']['result'] = harm_data.get('result', harm_data.get('status', 'N/A'))
            else:
                summary['harmonics']['result'] = 'PASS' if harm_data else 'FAIL'

        # 전체 compliance 판정
        results = [
            summary['frequency']['result'],
            summary['voltage']['result'],
            summary['thd']['result'],
            summary['unbalance']['result'],
            summary['flicker']['result'],
            summary['harmonics']['result']
        ]

        if 'FAIL' in results:
            summary['compliance'] = 'FAIL'
        elif all(r == 'PASS' for r in results):
            summary['compliance'] = 'PASS'
        else:
            summary['compliance'] = 'PARTIAL'

    except Exception as e:
        logging.error(f"EN50160 데이터 변환 실패: {e}")

    return summary

def load_diagnosis_parquet(channel: str, date_str: str, asset_name: str) -> Optional[Dict]:
    """
    진단 Parquet 파일 로드

    파일명: diagnosis_report_{asset_name}_{date}.parquet
    """
    try:
        filepath = REPORTS_DIR / channel / f"diagnosis_report_{asset_name}_{date_str}.parquet"

        if not filepath.exists():
            logging.warning(f"진단 파일 없음: {filepath}")
            return None

        df = pd.read_parquet(filepath)
        if df.empty:
            return None

        record = df.iloc[0].to_dict()

        result = {
            "asset_name": record.get("asset_name"),
            "channel_name": record.get("channel_name"),
            "diagnosis": None,
            "powerquality": None
        }

        # diagnosis 데이터 파싱
        if record.get("diagnosis"):
            diag_data = parse_json_field(record["diagnosis"])
            if diag_data:
                result["diagnosis"] = {
                    "timestamp": diag_data.get("timestamp"),
                    "main": parse_json_field(diag_data.get("main", [])),
                    "detail": parse_json_field(diag_data.get("detail", [])),
                    "trends": parse_json_field(diag_data.get("trends"))
                }

        # powerquality 데이터 파싱
        if record.get("powerquality"):
            pq_data = parse_json_field(record["powerquality"])
            if pq_data:
                result["powerquality"] = {
                    "timestamp": pq_data.get("timestamp"),
                    "main": parse_json_field(pq_data.get("main", [])),
                    "detail": parse_json_field(pq_data.get("detail", [])),
                    "trends": parse_json_field(pq_data.get("trends"))
                }

        logging.info(f"진단 데이터 로드 완료: {filepath.name}")
        return convert_to_serializable(result)

    except Exception as e:
        logging.error(f"진단 Parquet 로드 실패: {e}")
        return None


def load_en50160_parquet(channel: str, date_str: str) -> Optional[Dict]:
    """EN50160 주간 Parquet 파일 로드 (EN50160ReportProcessor 사용)"""
    try:
        filename = f"en50160_weekly_{date_str}.parquet"
        filepath = REPORTS_DIR / channel / filename

        if not filepath.exists():
            logging.warning(f"EN50160 파일 없음: {filepath}")
            return None

        # EN50160ReportProcessor 사용 (timeseries, histogram 포함)
        try:
            from routes.api import get_setupContext

            config = WeeklyReportConfig(output_dir=str(REPORTS_DIR))
            local_processor = EN50160ReportProcessor(config)

            ch_key = 'main' if channel.lower() == 'main' else 'sub'
            setupdict = get_setupContext()
            if setupdict and ch_key in setupdict:
                ch_setup = setupdict[ch_key]
                pt_info = ch_setup.get('ptInfo', {})
                local_processor.set_limits(
                    nominal_voltage=pt_info.get('vnorminal'),
                    nominal_frequency=pt_info.get('fnominal', 60.0)
                )

            result = local_processor.get_all_chart_data(channel, filename)
            if result:
                logging.info(f"EN50160 데이터 로드 완료 (Processor): {filename}")
                return convert_to_serializable(result)

        except ImportError as e:
            logging.warning(f"EN50160ReportProcessor import 실패, 기본 로직 사용: {e}")

        # 기본 로직 (fallback) - 기존 코드 유지
        table = pq.read_table(filepath)
        summary = None
        if table.schema.metadata and b'en50160_summary' in table.schema.metadata:
            summary_json = table.schema.metadata[b'en50160_summary'].decode('utf-8')
            summary = json.loads(summary_json)

        df = table.to_pandas()
        if df.empty:
            return None

        result = {
            "summary": summary,
            "frequency": _get_frequency_stats(df),
            "voltage": _get_voltage_stats(df),
            "thd": _get_thd_stats(df),
            "unbalance": _get_unbalance_stats(df),
            "flicker": _get_flicker_stats(df),
            "harmonics": _get_harmonics_stats(df),
            "total_samples": len(df)
        }
        logging.info(f"EN50160 데이터 로드 완료 (기본): {filename}")
        return convert_to_serializable(result)

    except Exception as e:
        logging.error(f"EN50160 Parquet 로드 실패: {e}")
        import traceback
        traceback.print_exc()
        return None

def _get_frequency_stats(df: pd.DataFrame) -> Optional[Dict]:
    """주파수 통계 계산"""
    col = "frequency_avg"
    if col not in df.columns:
        return None

    values = df[col].dropna().values
    if len(values) == 0:
        return None

    nominal = 60.0  # 기본값
    limit_99_5_min = nominal * 0.99
    limit_99_5_max = nominal * 1.01

    in_range = np.sum((values >= limit_99_5_min) & (values <= limit_99_5_max)) / len(values) * 100

    return {
        "statistics": {
            "min": float(np.min(values)),
            "max": float(np.max(values)),
            "avg": float(np.mean(values)),
            "in_range_99_5_percent": round(in_range, 2),
            "result_99_5": "PASS" if in_range >= 99.5 else "FAIL"
        },
        "limits": {
            "nominal": nominal,
            "limit_99_5": {"min": limit_99_5_min, "max": limit_99_5_max}
        }
    }


def _get_voltage_stats(df: pd.DataFrame) -> Optional[Dict]:
    """전압 통계 계산"""
    phases = {}
    nominal = 22900.0  # 기본값
    limit_95_min = nominal * 0.9
    limit_95_max = nominal * 1.1

    for phase in range(1, 4):
        col = f"voltage_l{phase}"
        if col not in df.columns:
            continue

        values = df[col].dropna().values
        if len(values) == 0:
            continue

        in_range = np.sum((values >= limit_95_min) & (values <= limit_95_max)) / len(values) * 100

        phases[f"L{phase}"] = {
            "statistics": {
                "min": float(np.min(values)),
                "max": float(np.max(values)),
                "avg": float(np.mean(values)),
                "in_range_95_percent": round(in_range, 2),
                "result_95": "PASS" if in_range >= 95.0 else "FAIL"
            }
        }

    if not phases:
        return None

    return {
        "phases": phases,
        "limits": {
            "nominal": nominal,
            "limit_95": {"min": limit_95_min, "max": limit_95_max}
        }
    }


def _get_thd_stats(df: pd.DataFrame) -> Optional[Dict]:
    """THD 통계 계산"""
    phases = {}
    limit_95 = 8.0

    for phase in range(1, 4):
        col = f"voltage_thd_l{phase}"
        if col not in df.columns:
            continue

        values = df[col].dropna().values
        if len(values) == 0:
            continue

        in_range = np.sum(values <= limit_95) / len(values) * 100

        phases[f"L{phase}"] = {
            "statistics": {
                "max": float(np.max(values)),
                "avg": float(np.mean(values)),
                "percentile_95": float(np.percentile(values, 95)),
                "in_range_95_percent": round(in_range, 2),
                "result": "PASS" if in_range >= 95.0 else "FAIL"
            }
        }

    if not phases:
        return None

    return {
        "phases": phases,
        "limits": {"limit_95": limit_95}
    }


def _get_unbalance_stats(df: pd.DataFrame) -> Optional[Dict]:
    """전압 불평형 통계 계산"""
    col = "voltage_unbalance_0"
    if col not in df.columns:
        return None

    values = df[col].dropna().values
    if len(values) == 0:
        return None

    limit_95 = 2.0
    in_range = np.sum(values <= limit_95) / len(values) * 100

    return {
        "statistics": {
            "max": float(np.max(values)),
            "avg": float(np.mean(values)),
            "percentile_95": float(np.percentile(values, 95)),
            "in_range_95_percent": round(in_range, 2),
            "result": "PASS" if in_range >= 95.0 else "FAIL"
        },
        "limits": {"limit_95": limit_95}
    }


def _get_flicker_stats(df: pd.DataFrame) -> Optional[Dict]:
    """플리커 통계 계산"""
    result = {"pst": {}, "plt": {}}
    plt_limit = 1.0

    for phase in range(1, 4):
        # Pst
        pst_col = f"pst_l{phase}"
        if pst_col in df.columns:
            values = df[pst_col].dropna().values
            if len(values) > 0:
                result["pst"][f"L{phase}"] = {
                    "statistics": {
                        "max": float(np.max(values)),
                        "avg": float(np.mean(values)),
                        "percentile_95": float(np.percentile(values, 95))
                    }
                }

        # Plt
        plt_col = f"plt_l{phase}"
        if plt_col in df.columns:
            values = df[plt_col].dropna().values
            if len(values) > 0:
                in_range = np.sum(values <= plt_limit) / len(values) * 100
                result["plt"][f"L{phase}"] = {
                    "statistics": {
                        "max": float(np.max(values)),
                        "avg": float(np.mean(values)),
                        "percentile_95": float(np.percentile(values, 95)),
                        "in_range_95_percent": round(in_range, 2),
                        "result": "PASS" if in_range >= 95.0 else "FAIL"
                    }
                }

    if not result["pst"] and not result["plt"]:
        return None

    result["limits"] = {"plt_95": plt_limit}
    return result


def _get_harmonics_stats(df: pd.DataFrame) -> Optional[Dict]:
    """고조파 통계 계산"""
    phases = {}

    harmonics_limits = {
        "h2": 2.0, "h3": 5.0, "h4": 1.0, "h5": 6.0,
        "h6": 0.5, "h7": 5.0, "h8": 0.5, "h9": 1.5,
        "h10": 0.5, "h11": 3.5, "h12": 0.5, "h13": 3.0,
    }

    for phase in range(1, 4):
        col = f"harmonics_l{phase}"
        if col not in df.columns:
            continue

        # harmonics_l1은 리스트 컬럼
        try:
            all_harmonics = np.array(df[col].tolist())
            if all_harmonics.shape[0] == 0:
                continue

            phase_data = {}
            for h in range(min(12, all_harmonics.shape[1])):  # H2 ~ H13
                h_num = h + 2
                h_key = f"h{h_num}"
                h_values = all_harmonics[:, h]
                limit = harmonics_limits.get(h_key, 0.5)

                in_range = np.sum(h_values <= limit) / len(h_values) * 100

                phase_data[h_key] = {
                    "max": float(np.max(h_values)),
                    "avg": float(np.mean(h_values)),
                    "limit": limit,
                    "in_range_95_percent": round(in_range, 2),
                    "result": "PASS" if in_range >= 95.0 else "FAIL"
                }

            phases[f"L{phase}"] = phase_data
        except Exception as e:
            logging.warning(f"고조파 파싱 실패 L{phase}: {e}")
            continue

    if not phases:
        return None

    return {
        "phases": phases,
        "limits": harmonics_limits
    }


def load_energy_data(channel: str, start_date: datetime, end_date: datetime) -> Optional[Dict]:
    """전력량 데이터 조회 - 일별/시간별/월별"""
    try:
        from .api import get_dailyEnergy, get_hourEnergy, get_monthlyEnergy

        daily_data = []
        hourly_data = []
        monthly_data = []
        hourly_by_hour = {h: [] for h in range(24)}

        # 월별 데이터 조회
        year = start_date.strftime("%Y")
        try:
            monthly_response = get_monthlyEnergy(channel, year)
            if monthly_response.get('success') and monthly_response.get('monthlyData'):
                monthly_data = monthly_response['monthlyData']
        except Exception as e:
            logging.warning(f"월별 데이터 조회 실패: {e}")

        current_date = start_date
        processed_months = set()

        while current_date <= end_date:
            date_str = current_date.strftime("%Y-%m-%d")
            year_month = current_date.strftime("%Y-%m")

            if year_month not in processed_months:
                daily_response = get_dailyEnergy(channel, year_month)
                if daily_response.get('success') and daily_response.get('dailyData'):
                    for item in daily_response['dailyData']:
                        item_date = item.get('date', '')
                        if start_date.strftime("%Y-%m-%d") <= item_date <= end_date.strftime("%Y-%m-%d"):
                            if not any(d['date'] == item_date for d in daily_data):
                                daily_data.append({'date': item_date, 'value': item.get('value', 0)})
                processed_months.add(year_month)

            hourly_response = get_hourEnergy(channel, date_str)
            if hourly_response.get('success') and hourly_response.get('hourlyData'):
                for item in hourly_response['hourlyData']:
                    hour = item.get('hour', 0)
                    value = item.get('value', 0)
                    hourly_by_hour[hour].append(value)
                    if current_date.date() == end_date.date():
                        hourly_data.append({'hour': hour, 'value': value})

            current_date += timedelta(days=1)

        if not daily_data:
            return None

        daily_data.sort(key=lambda x: x['date'])
        hourly_data.sort(key=lambda x: x.get('hour', 0))

        hourly_avg = {h: sum(v) / len(v) if v else 0 for h, v in hourly_by_hour.items()}
        peak_hour = max(hourly_avg, key=hourly_avg.get)
        off_peak_hour = min(hourly_avg, key=hourly_avg.get)

        return {
            'daily': daily_data,
            'hourly': hourly_data,
            'monthly': monthly_data,
            'hourlyPattern': {
                'peakHour': {'hour': peak_hour, 'avgValue': hourly_avg[peak_hour]},
                'offPeakHour': {'hour': off_peak_hour, 'avgValue': hourly_avg[off_peak_hour]}
            }
        }
    except Exception as e:
        logging.error(f"전력량 데이터 조회 실패: {e}")
        return None

def get_asset_info_from_redis(channel: str, redis_state=None) -> Optional[Dict]:
    """
    Redis에서 설비 정보 조회
    """
    try:
        if redis_state is None or redis_state.client is None:
            return None

        # ChannelData에서 정보 조회
        if redis_state.client.hexists("Equipment", "ChannelData"):
            ch_data = json.loads(redis_state.client.hget("Equipment", "ChannelData"))
            ch_key = "main" if channel.lower() == "main" else "sub"

            if ch_key in ch_data:
                return {
                    "ratedVoltage": ch_data[ch_key].get("RatedVoltage"),
                    "ratedCurrent": ch_data[ch_key].get("RatedCurrent"),
                    "ratedFrequency": ch_data[ch_key].get("RatedFrequency")
                }

        # System setup에서 정보 조회
        setup_json = redis_state.client.hget("System", "setup")
        if setup_json:
            setup = json.loads(setup_json)
            for ch in setup.get("channel", []):
                if ch.get("channel") == channel:
                    asset_info = ch.get("assetInfo", {})
                    pt_info = ch.get(channel, {}).get("ptInfo", {})
                    ct_info = ch.get(channel, {}).get("ctInfo", {})

                    return {
                        "name": asset_info.get("name"),
                        "type": asset_info.get("type"),
                        "ratedVoltage": pt_info.get("vnorminal"),
                        "ratedCurrent": ct_info.get("inorminal"),
                        "ratedFrequency": pt_info.get("linefrequency", 60)
                    }

        return None

    except Exception as e:
        logging.error(f"Redis 설비 정보 조회 실패: {e}")
        return None


# ============================================
# API 라우터
# ============================================
@router.get("/downloadWeeklyReport/{channel}/{date}")
async def download_weekly_report(
        channel: str,
        date: str,
        locale: str = Query(default='en', regex='^(ko|en)$'),
        asset_name: str = Query(default=None)
):
    try:
        logging.info(f"📄 통합 주간 리포트 생성 요청: channel={channel}, date={date}, locale={locale}")

        # 주간 범위 계산
        start_date, end_date = get_week_range_from_date(date)
        report_period = f"{start_date.strftime('%Y-%m-%d')} ~ {end_date.strftime('%Y-%m-%d')}"

        # 설비 정보 조회
        # 실제 환경에서는 redis_state를 import해서 사용
        # asset_info_from_redis = get_asset_info_from_redis(channel, redis_state)
        asset_info = {
            "name": asset_name or "Unknown",
            "type": "Motor",
            "channel": channel,
            "ratedVoltage": 380,
            "ratedCurrent": 100,
            "ratedFrequency": 60
        }

        # Redis에서 정보 조회 시도
        # if not asset_name:
        #     redis_info = get_asset_info_from_redis(channel, redis_state)
        #     if redis_info:
        #         asset_info.update(redis_info)

        if asset_name:
            asset_info["name"] = asset_name

        # 진단 데이터 로드
        diagnosis_result = load_diagnosis_parquet(channel, date, asset_info["name"])

        diagnosis_data = None
        powerquality_data = None

        if diagnosis_result:
            diagnosis_data = diagnosis_result.get("diagnosis")
            powerquality_data = diagnosis_result.get("powerquality")

        # EN50160 데이터 로드
        en50160_data = load_en50160_parquet(channel, date)

        if en50160_data is None or en50160_data.get('summary') is None:
            en50160_data = get_en50160_from_redis(channel, en50160_data)

        # 전력량 데이터 로드
        # 실제 환경에서는 influx_state를 import해서 사용
        energy_data = load_energy_data(channel, start_date, end_date)
        # energy_data = None  # 실제 환경에서 활성화

        itic_events = get_itic_data_for_report(channel)
        # 출력 파일 경로
        temp_dir = tempfile.gettempdir()
        report_id = str(uuid.uuid4())[:8]
        output_path = f'{temp_dir}/weekly_report_{report_id}.docx'

        # 리포트 생성
        generate_weekly_report(
            asset_info=asset_info,
            diagnosis_data=diagnosis_data,
            powerquality_data=powerquality_data,
            en50160_data=en50160_data,
            energy_data=energy_data,
            itic_events=itic_events,  # 추가됨
            report_period=report_period,
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            locale=locale,
            output_path=output_path
        )

        if not Path(output_path).exists():
            raise HTTPException(status_code=500, detail="Report file not created")

        # 파일명 생성
        filename = f"weekly_report_{channel}_{date}.docx"

        logging.info(f"✅ 통합 주간 리포트 생성 완료: {filename}")

        return FileResponse(
            path=output_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"❌ 리포트 생성 실패: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/weeklyReportData/{channel}/{date}")
async def get_weekly_report_data(
        channel: str,
        date: str,
        asset_name: str = Query(default=None)
):
    """
    통합 주간 리포트 데이터 조회 (JSON)

    미리보기나 디버깅용
    """
    try:
        start_date, end_date = get_week_range_from_date(date)

        # 진단 데이터 로드
        diagnosis_result = load_diagnosis_parquet(channel, date, asset_name or "Unknown")

        # EN50160 데이터 로드
        en50160_data = load_en50160_parquet(channel, date)

        return {
            "success": True,
            "data": {
                "channel": channel,
                "date": date,
                "reportPeriod": f"{start_date.strftime('%Y-%m-%d')} ~ {end_date.strftime('%Y-%m-%d')}",
                "diagnosis": diagnosis_result.get("diagnosis") if diagnosis_result else None,
                "powerquality": diagnosis_result.get("powerquality") if diagnosis_result else None,
                "en50160": en50160_data
            }
        }

    except Exception as e:
        logging.error(f"리포트 데이터 조회 실패: {e}")
        return {"success": False, "msg": str(e)}


# @router.get("/weeklyReportDates/{channel}")
# async def get_weekly_report_dates(channel: str):
#     """
#     사용 가능한 주간 리포트 날짜 목록 조회
#     """
#     try:
#         channel_dir = REPORTS_DIR / channel
#
#         if not channel_dir.exists():
#             return {"success": True, "data": []}
#
#         dates = set()
#
#         # EN50160 파일에서 날짜 추출
#         for file in channel_dir.glob("en50160_weekly_*.parquet"):
#             try:
#                 parts = file.stem.split("_")
#                 date_str = parts[-1]
#                 if len(date_str) == 8 and date_str.isdigit():
#                     dates.add(date_str)
#             except:
#                 continue
#
#         # 진단 파일에서 날짜 추출
#         for file in channel_dir.glob("diagnosis_report_*.parquet"):
#             try:
#                 parts = file.stem.split("_")
#                 date_str = parts[-1]
#                 if len(date_str) == 8 and date_str.isdigit():
#                     dates.add(date_str)
#             except:
#                 continue
#
#         sorted_dates = sorted(list(dates), reverse=True)
#
#         return {"success": True, "data": sorted_dates}
#
#     except Exception as e:
#         logging.error(f"날짜 목록 조회 실패: {e}")
#         return {"success": False, "msg": str(e)}