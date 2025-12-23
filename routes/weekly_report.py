"""
통합 주간 리포트 생성기 (Word)

설비진단 + PQ진단 + EN50160 + 전력량을 통합한 Word 리포트 생성
python-docx 사용
"""

import json
import logging
import os
import tempfile
import uuid
import warnings

from typing import Dict, Any, List, Optional, Tuple

import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# matplotlib 기본 설정 (모듈 로드 시 바로 적용)
plt.rcParams['axes.unicode_minus'] = False

# matplotlib 폰트 관련 경고 억제
warnings.filterwarnings('ignore', category=UserWarning, module='matplotlib')
warnings.filterwarnings('ignore', message='.*glyph.*')
warnings.filterwarnings('ignore', message='.*Substituting.*')

# matplotlib 로거 레벨 조정
logging.getLogger('matplotlib').setLevel(logging.DEBUG)
logging.getLogger('matplotlib.font_manager').setLevel(logging.DEBUG)

logger = logging.getLogger(__name__)

# ============================================
# 색상 상수
# ============================================
COLORS = {
    'PRIMARY': RGBColor(30, 64, 175),      # 진한 파랑
    'SECONDARY': RGBColor(100, 116, 139),  # 슬레이트 그레이
    'SUCCESS': RGBColor(16, 185, 129),     # 초록
    'WARNING': RGBColor(234, 179, 8),      # 노랑
    'DANGER': RGBColor(239, 68, 68),       # 빨강
    'ORANGE': RGBColor(249, 115, 22),      # 주황
    'WHITE': RGBColor(255, 255, 255),
    'BLACK': RGBColor(31, 41, 55),
    'LIGHT_GRAY': RGBColor(243, 244, 246),
    'HEADER_BG': '1E3A5F',
}

STATUS_CONFIG = {
    'diagnosis': {
        0: {'bg': '64748B', 'text_en': 'Stop', 'text_ko': '정지', 'color': RGBColor(100, 116, 139)},
        1: {'bg': '10B981', 'text_en': 'OK', 'text_ko': '정상', 'color': RGBColor(16, 185, 129)},
        2: {'bg': 'EAB308', 'text_en': 'Warning', 'text_ko': '주의', 'color': RGBColor(234, 179, 8)},
        3: {'bg': 'F97316', 'text_en': 'Inspect', 'text_ko': '경고', 'color': RGBColor(249, 115, 22)},
        4: {'bg': 'EF4444', 'text_en': 'Repair', 'text_ko': '위험', 'color': RGBColor(239, 68, 68)},
    },
    'powerquality': {
        0: {'bg': '64748B', 'text_en': 'Stop', 'text_ko': '정지', 'color': RGBColor(100, 116, 139)},
        1: {'bg': '10B981', 'text_en': 'OK', 'text_ko': '정상', 'color': RGBColor(16, 185, 129)},
        2: {'bg': 'EAB308', 'text_en': 'Warning', 'text_ko': '주의', 'color': RGBColor(234, 179, 8)},
        3: {'bg': 'F97316', 'text_en': 'Inspect', 'text_ko': '경고', 'color': RGBColor(249, 115, 22)},
        4: {'bg': 'EF4444', 'text_en': 'Repair', 'text_ko': '위험', 'color': RGBColor(239, 68, 68)},
    }
}


# ============================================
# 유틸리티 함수
# ============================================
def get_text(locale: str, ko: str, en: str) -> str:
    """로케일에 따른 텍스트 반환"""
    return ko if locale == 'ko' else en


def format_number(value, decimals: int = 2) -> str:
    """숫자 포맷팅"""
    if value is None or (isinstance(value, float) and np.isnan(value)):
        return '-'
    try:
        return f"{float(value):.{decimals}f}"
    except:
        return str(value)


def set_cell_shading(cell, color: str):
    """셀 배경색 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    """셀 여백 설정"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(margin_value))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


# 폰트 캐시 (한 번만 초기화)
_font_prop_cache = None
_font_initialized = False


def setup_korean_font():
    """한글 폰트 설정 (한 번만 초기화)"""
    global _font_prop_cache, _font_initialized

    # 이미 초기화되었으면 캐시된 값 반환
    if _font_initialized:
        return _font_prop_cache

    font_paths = [
        '/home/root/webserver/fonts/NanumGothicCoding.ttf',
        '/home/root/webserver/fonts/NanumGothicCoding-Bold.ttf',
    ]

    for font_path in font_paths:
        if os.path.exists(font_path):
            try:
                fm.fontManager.addfont(font_path)
                font_prop = fm.FontProperties(fname=font_path)
                font_name = font_prop.get_name()
                plt.rcParams['font.family'] = font_name
                plt.rcParams['axes.unicode_minus'] = False
                _font_prop_cache = font_prop
                _font_initialized = True
                return font_prop
            except Exception as e:
                logger.debug(f"폰트 로드 실패: {font_path}, {e}")
                continue

    # 폰트 없으면 기본 설정만
    plt.rcParams['axes.unicode_minus'] = False
    _font_initialized = True
    return None


# ============================================
# 테이블 생성 헬퍼
# ============================================
def create_info_table(doc, rows: List[Tuple[str, str]], col_widths: Tuple[float, float] = (4, 10)):
    """정보 표시용 2열 테이블 생성"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'

    for i, (label, value) in enumerate(rows):
        # 라벨 셀
        label_cell = table.rows[i].cells[0]
        label_cell.width = Cm(col_widths[0])
        set_cell_shading(label_cell, 'F3F4F6')
        p = label_cell.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.bold = True

        # 값 셀
        value_cell = table.rows[i].cells[1]
        value_cell.width = Cm(col_widths[1])
        p = value_cell.paragraphs[0]
        run = p.add_run(str(value) if value else '-')
        run.font.size = Pt(10)

    return table


def create_data_table(doc, headers: List[str], data: List[List[Any]],
                      col_widths: List[float] = None,
                      header_bg: str = '1E3A5F'):
    """데이터 테이블 생성"""
    if not data:
        return None

    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'

    # 헤더 행
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        if col_widths and i < len(col_widths):
            cell.width = Cm(col_widths[i])
        set_cell_shading(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # 데이터 행
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            if col_widths and col_idx < len(col_widths):
                cell.width = Cm(col_widths[col_idx])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 딕셔너리면 특수 처리 (배경색 등)
            if isinstance(value, dict):
                run = p.add_run(str(value.get('text', '')))
                run.font.size = Pt(10)
                if value.get('bold'):
                    run.font.bold = True
                if value.get('color'):
                    run.font.color.rgb = value['color']
                if value.get('bg'):
                    set_cell_shading(cell, value['bg'])
            else:
                run = p.add_run(str(value) if value is not None else '-')
                run.font.size = Pt(10)

    return table


# ============================================
# 차트 생성 함수
# ============================================
def create_diagnosis_bar_chart(chart_path: str, main_data: List[Dict],
                                mode: str = 'diagnosis', locale: str = 'en') -> str:
    """진단 결과 바차트 생성"""
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(10, max(4, len(main_data) * 0.5)))

    labels = []
    values = []
    colors = []

    color_map = {
        0: '#64748B',
        1: '#10B981',
        2: '#EAB308',
        3: '#F97316',
        4: '#EF4444',
    }

    for item in main_data:
        if locale == 'ko':
            label = item.get('title_ko') or item.get('title') or item.get('item_name', '')
        else:
            label = item.get('title_en') or item.get('title') or item.get('item_name', '')
        labels.append(label)

        status = item.get('status', 0)
        values.append(status if status > 0 else 0.2)
        colors.append(color_map.get(int(status), color_map[0]))

    y_pos = np.arange(len(labels))
    ax.barh(y_pos, values, color=colors, height=0.6, edgecolor='white', linewidth=1)

    ax.set_yticks(y_pos)
    if font_prop:
        ax.set_yticklabels(labels, fontsize=10, fontproperties=font_prop)
    else:
        ax.set_yticklabels(labels, fontsize=10)

    ax.set_xlim(0, 4.5)
    ax.set_xticks([0, 1, 2, 3, 4])

    if mode == 'powerquality':
        status_labels = ['정지', '정상', '주의', '경고', '위험'] if locale == 'ko' else ['Stop', 'OK', 'Warning', 'Inspect', 'Danger']
    else:
        status_labels = ['정지', '정상', '주의', '경고', '위험'] if locale == 'ko' else ['Stop', 'OK', 'Warning', 'Inspect', 'Repair']

    if font_prop:
        ax.set_xticklabels(status_labels, fontsize=9, fontproperties=font_prop)
    else:
        ax.set_xticklabels(status_labels, fontsize=9)

    ax.xaxis.grid(True, linestyle='--', alpha=0.3)
    ax.set_axisbelow(True)

    title = get_text(locale, '설비 진단 현황' if mode == 'diagnosis' else '전력품질 진단 현황',
                     'Diagnostic Status' if mode == 'diagnosis' else 'Power Quality Status')
    if font_prop:
        ax.set_title(title, fontsize=12, fontweight='bold', pad=10, fontproperties=font_prop)
    else:
        ax.set_title(title, fontsize=12, fontweight='bold', pad=10)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def create_trend_chart(chart_path: str, trend_data: Dict, locale: str = 'en', is_status_chart: bool = True) -> Optional[str]:
    """
    트렌드 차트 생성

    Args:
        chart_path: 저장 경로
        trend_data: 차트 데이터
        locale: 언어
        is_status_chart: True면 Y축을 상태 라벨(정지/정상/주의/경고/위험)로 표시
    """
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(10, 4))

    labels = trend_data.get('lineLabels', [])
    datasets = trend_data.get('lineData', [])
    chart_title = trend_data.get('lineTitle', 'Trend')

    if not labels or not datasets:
        plt.close()
        return None

    x = np.arange(len(labels))

    # 상태별 색상 매핑
    status_colors = {
        0: '#64748B',  # 정지 - 회색
        1: '#10B981',  # 정상 - 초록
        2: '#EAB308',  # 주의 - 노랑
        3: '#F97316',  # 경고 - 주황
        4: '#EF4444',  # 위험 - 빨강
    }

    for dataset in datasets:
        name = dataset.get('name', '')
        data = dataset.get('data', [])
        is_threshold = dataset.get('isThreshold', False)

        if is_threshold:
            ax.plot(x, data, '--', label=name, alpha=0.7, linewidth=1.5)
        else:
            # 상태 차트면 색상으로 구분
            if is_status_chart and name == 'Status':
                # 각 포인트를 상태별 색상으로 표시
                for i, (xi, yi) in enumerate(zip(x, data)):
                    color = status_colors.get(int(yi), '#64748B')
                    ax.scatter(xi, yi, c=color, s=50, zorder=5)
                # 라인도 그리기
                ax.plot(x, data, '-', color='#94A3B8', linewidth=1, alpha=0.5)
            else:
                ax.plot(x, data, '-', label=name, linewidth=2)

    # X축 라벨 포맷팅
    def format_xaxis_label(label):
        if 'T' in str(label):
            return str(label).split('T')[1][:5] if 'T' in str(label) else str(label)[-5:]
        return str(label)[-5:] if len(str(label)) > 5 else str(label)

    if len(labels) > 10:
        step = len(labels) // 10
        ax.set_xticks(x[::step])
        ax.set_xticklabels([format_xaxis_label(l) for l in labels[::step]], rotation=45, ha='right', fontsize=8)
    else:
        ax.set_xticks(x)
        ax.set_xticklabels([format_xaxis_label(l) for l in labels], rotation=45, ha='right', fontsize=8)

    # Y축 상태 라벨 설정 (상태 차트인 경우)
    if is_status_chart:
        ax.set_ylim(-0.5, 4.5)
        ax.set_yticks([0, 1, 2, 3, 4])
        if locale == 'ko':
            status_labels = ['정지', '정상', '주의', '경고', '위험']
        else:
            status_labels = ['Stop', 'OK', 'Warning', 'Inspect', 'Danger']

        if font_prop:
            ax.set_yticklabels(status_labels, fontproperties=font_prop)
        else:
            ax.set_yticklabels(status_labels)

    if font_prop:
        ax.set_title(chart_title, fontsize=12, fontweight='bold', pad=10, fontproperties=font_prop)
    else:
        ax.set_title(chart_title, fontsize=12, fontweight='bold', pad=10)

    xlabel = get_text(locale, '시간', 'Time')
    ylabel = get_text(locale, '상태', 'Status') if is_status_chart else get_text(locale, '값', 'Value')

    if font_prop:
        ax.set_xlabel(xlabel, fontsize=10, fontproperties=font_prop)
        ax.set_ylabel(ylabel, fontsize=10, fontproperties=font_prop)
    else:
        ax.set_xlabel(xlabel, fontsize=10)
        ax.set_ylabel(ylabel, fontsize=10)

    ax.grid(True, linestyle='--', alpha=0.3)

    # 상태 차트가 아닐 때만 범례 표시
    if not is_status_chart:
        ax.legend(loc='upper right', fontsize=8)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def create_energy_chart(chart_path: str, daily_data: List[Dict], locale: str = 'en') -> str:
    """일별 전력량 바차트 생성"""
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(10, 5))

    dates = [d.get('date', '')[-5:] for d in daily_data]  # MM-DD 형식
    values = [d.get('value', 0) for d in daily_data]

    bars = ax.bar(dates, values, color='#3B82F6', edgecolor='white', linewidth=1)

    # 최대값 강조
    if values and max(values) > 0:
        max_idx = values.index(max(values))
        if bars:
            bars[max_idx].set_color('#EF4444')

    title = get_text(locale, '일별 전력량', 'Daily Energy Consumption')
    if font_prop:
        ax.set_title(title, fontsize=12, fontweight='bold', fontproperties=font_prop)
        ax.set_xlabel(get_text(locale, '날짜', 'Date'), fontsize=10, fontproperties=font_prop)
        ax.set_ylabel(get_text(locale, '전력량 (kWh)', 'Energy (kWh)'), fontsize=10, fontproperties=font_prop)
    else:
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_xlabel(get_text(locale, '날짜', 'Date'), fontsize=10)
        ax.set_ylabel(get_text(locale, '전력량 (kWh)', 'Energy (kWh)'), fontsize=10)

    plt.xticks(rotation=45, ha='right', fontsize=8)
    ax.yaxis.grid(True, linestyle='--', alpha=0.3)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def create_hourly_energy_chart(chart_path: str, hourly_data: List[Dict], locale: str = 'en') -> str:
    """시간별 전력량 차트 생성"""
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(10, 4))

    hours = [d.get('hour', i) for i, d in enumerate(hourly_data)]
    values = [d.get('value', 0) for d in hourly_data]

    # 24시간 데이터 보장
    if len(hours) < 24:
        full_hours = list(range(24))
        full_values = [0] * 24
        for h, v in zip(hours, values):
            if 0 <= h < 24:
                full_values[h] = v
        hours = full_hours
        values = full_values

    bars = ax.bar(hours, values, color='#10B981', edgecolor='white', linewidth=0.5)

    # 피크 시간 강조
    if values and max(values) > 0:
        max_idx = values.index(max(values))
        bars[max_idx].set_color('#EF4444')

    title = get_text(locale, '시간별 전력량', 'Hourly Energy Consumption')
    if font_prop:
        ax.set_title(title, fontsize=12, fontweight='bold', fontproperties=font_prop)
        ax.set_xlabel(get_text(locale, '시간', 'Hour'), fontsize=10, fontproperties=font_prop)
        ax.set_ylabel(get_text(locale, '전력량 (kWh)', 'Energy (kWh)'), fontsize=10, fontproperties=font_prop)
    else:
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_xlabel(get_text(locale, '시간', 'Hour'), fontsize=10)
        ax.set_ylabel(get_text(locale, '전력량 (kWh)', 'Energy (kWh)'), fontsize=10)

    ax.set_xticks(range(0, 24, 2))
    ax.set_xticklabels([f'{h}:00' for h in range(0, 24, 2)], fontsize=8)
    ax.yaxis.grid(True, linestyle='--', alpha=0.3)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def create_monthly_energy_chart(chart_path: str, monthly_data: List[Dict], locale: str = 'en') -> str:
    """월별 전력량 차트 생성"""
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(10, 4))

    months = [d.get('month', d.get('year_month', '')) for d in monthly_data]
    values = [d.get('value', 0) for d in monthly_data]

    # 월 라벨 포맷팅
    if locale == 'ko':
        month_labels = [f'{m}월' if isinstance(m, int) else str(m) for m in months]
    else:
        month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
        month_labels = [month_names[m-1] if isinstance(m, int) and 1 <= m <= 12 else str(m) for m in months]

    bars = ax.bar(range(len(months)), values, color='#8B5CF6', edgecolor='white', linewidth=1)

    # 최대값 강조
    if values and max(values) > 0:
        max_idx = values.index(max(values))
        bars[max_idx].set_color('#EF4444')

    title = get_text(locale, '월별 전력량', 'Monthly Energy Consumption')
    if font_prop:
        ax.set_title(title, fontsize=12, fontweight='bold', fontproperties=font_prop)
        ax.set_xlabel(get_text(locale, '월', 'Month'), fontsize=10, fontproperties=font_prop)
        ax.set_ylabel(get_text(locale, '전력량 (kWh)', 'Energy (kWh)'), fontsize=10, fontproperties=font_prop)
    else:
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_xlabel(get_text(locale, '월', 'Month'), fontsize=10)
        ax.set_ylabel(get_text(locale, '전력량 (kWh)', 'Energy (kWh)'), fontsize=10)

    ax.set_xticks(range(len(months)))
    ax.set_xticklabels(month_labels, fontsize=9)
    ax.yaxis.grid(True, linestyle='--', alpha=0.3)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def create_en50160_timeseries_chart(chart_path: str, data: Dict, title: str,
                                     limits: Dict = None, locale: str = 'en') -> Optional[str]:
    """
    EN50160 시계열 차트 생성 (라인 차트 + 한계선)

    Args:
        data: {"labels": [...], "data": [...]}
        title: 차트 제목
        limits: {"limit_99_5": {"min": ..., "max": ...}, ...}
    """
    font_prop = setup_korean_font()

    labels = data.get('labels', [])
    values = data.get('data', [])

    if not labels or not values:
        return None

    fig, ax = plt.subplots(figsize=(10, 4))

    x = np.arange(len(labels))

    # 메인 데이터 라인
    ax.plot(x, values, '-', color='#0891B2', linewidth=1.5, label=title)

    # 한계선 그리기
    if limits:
        # 99.5% 한계 (노란색 점선)
        if 'limit_99_5' in limits:
            limit_99_5 = limits['limit_99_5']
            if isinstance(limit_99_5, dict):
                ax.axhline(y=limit_99_5.get('min'), color='#EAB308', linestyle='--', linewidth=1, alpha=0.8)
                ax.axhline(y=limit_99_5.get('max'), color='#EAB308', linestyle='--', linewidth=1, alpha=0.8)

        # 100% 한계 (빨간색 점선)
        if 'limit_100' in limits:
            limit_100 = limits['limit_100']
            if isinstance(limit_100, dict):
                ax.axhline(y=limit_100.get('min'), color='#EF4444', linestyle='--', linewidth=1, alpha=0.8)
                ax.axhline(y=limit_100.get('max'), color='#EF4444', linestyle='--', linewidth=1, alpha=0.8)

        # 단일 한계 (95%)
        if 'limit_95' in limits:
            limit_95 = limits['limit_95']
            if isinstance(limit_95, dict):
                ax.axhline(y=limit_95.get('min'), color='#EAB308', linestyle='--', linewidth=1, alpha=0.8)
                ax.axhline(y=limit_95.get('max'), color='#EAB308', linestyle='--', linewidth=1, alpha=0.8)
            elif isinstance(limit_95, (int, float)):
                ax.axhline(y=limit_95, color='#EAB308', linestyle='--', linewidth=1, alpha=0.8)

    # X축 라벨
    if len(labels) > 10:
        step = len(labels) // 10
        ax.set_xticks(x[::step])
        x_labels = [str(l)[-11:-6] if len(str(l)) > 11 else str(l) for l in labels[::step]]
        ax.set_xticklabels(x_labels, rotation=45, ha='right', fontsize=7)
    else:
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=7)

    if font_prop:
        ax.set_title(title, fontsize=11, fontweight='bold', fontproperties=font_prop)
    else:
        ax.set_title(title, fontsize=11, fontweight='bold')

    ax.grid(True, linestyle='--', alpha=0.3)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def create_en50160_histogram_chart(chart_path: str, histogram: Dict, title: str,
                                    stats: Dict = None, limits: Dict = None,
                                    locale: str = 'en') -> Optional[str]:
    """
    EN50160 히스토그램 차트 생성

    Args:
        histogram: {"bin_centers": [...], "counts": [...], "bins": [...]}
        title: 차트 제목
        stats: {"min": ..., "max": ...}
        limits: 한계값
    """
    font_prop = setup_korean_font()

    bin_centers = histogram.get('bin_centers', [])
    counts = histogram.get('counts', [])
    bins = histogram.get('bins', [])

    if not bin_centers or not counts:
        return None

    fig, ax = plt.subplots(figsize=(8, 3))

    # 바 차트
    bar_width = (bins[1] - bins[0]) * 0.8 if len(bins) > 1 else 0.1
    bars = ax.bar(bin_centers, counts, width=bar_width, color='#0891B2', edgecolor='white', linewidth=0.5)

    # 최소/최대 표시
    if stats:
        min_val = stats.get('min')
        max_val = stats.get('max')
        if min_val is not None:
            ax.axvline(x=min_val, color='#EF4444', linestyle='--', linewidth=1.5, alpha=0.8)
            ax.text(min_val, ax.get_ylim()[1] * 0.9, f'Min:{min_val}', fontsize=7, color='#EF4444', ha='center')
        if max_val is not None:
            ax.axvline(x=max_val, color='#EF4444', linestyle='--', linewidth=1.5, alpha=0.8)
            ax.text(max_val, ax.get_ylim()[1] * 0.9, f'Max:{max_val}', fontsize=7, color='#EF4444', ha='center')

    dist_title = get_text(locale, f'{title} 분포', f'{title} Distribution')
    if font_prop:
        ax.set_title(dist_title, fontsize=10, fontweight='bold', fontproperties=font_prop)
        ax.set_xlabel(title, fontsize=9, fontproperties=font_prop)
        ax.set_ylabel(get_text(locale, '빈도', 'Count'), fontsize=9, fontproperties=font_prop)
    else:
        ax.set_title(dist_title, fontsize=10, fontweight='bold')
        ax.set_xlabel(title, fontsize=9)
        ax.set_ylabel(get_text(locale, '빈도', 'Count'), fontsize=9)

    ax.yaxis.grid(True, linestyle='--', alpha=0.3)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def create_en50160_summary_chart(chart_path: str, summary: Dict, locale: str = 'en') -> str:
    """EN50160 요약 파이/바 차트 생성"""
    font_prop = setup_korean_font()

    items = [
        ('frequency', get_text(locale, '주파수', 'Frequency')),
        ('voltage', get_text(locale, '전압', 'Voltage')),
        ('thd', get_text(locale, 'THD', 'THD')),
        ('unbalance', get_text(locale, '불평형', 'Unbalance')),
        ('flicker', get_text(locale, '플리커', 'Flicker')),
        ('harmonics', get_text(locale, '고조파', 'Harmonics')),
    ]

    labels = []
    colors = []

    for key, label in items:
        labels.append(label)
        result = summary.get(key, {})
        if isinstance(result, dict):
            result = result.get('result', 'N/A')
        colors.append('#10B981' if result == 'PASS' else '#EF4444' if result == 'FAIL' else '#64748B')

    fig, ax = plt.subplots(figsize=(8, 4))

    y_pos = np.arange(len(labels))
    ax.barh(y_pos, [1] * len(labels), color=colors, height=0.6)

    ax.set_yticks(y_pos)
    if font_prop:
        ax.set_yticklabels(labels, fontsize=10, fontproperties=font_prop)
    else:
        ax.set_yticklabels(labels, fontsize=10)

    ax.set_xlim(0, 1.2)
    ax.set_xticks([])

    # 결과 텍스트 추가
    for i, (key, _) in enumerate(items):
        result = summary.get(key, {})
        if isinstance(result, dict):
            result = result.get('result', 'N/A')
        ax.text(0.5, i, result, ha='center', va='center', fontsize=11,
                fontweight='bold', color='white')

    title = get_text(locale, 'EN50160 판정 요약', 'EN50160 Compliance Summary')
    if font_prop:
        ax.set_title(title, fontsize=12, fontweight='bold', fontproperties=font_prop)
    else:
        ax.set_title(title, fontsize=12, fontweight='bold')

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


# ============================================
# 섹션 생성 함수
# ============================================
def add_cover_page(doc, asset_info: Dict, report_period: str, generated_at: str, locale: str):
    """표지 생성"""
    # 여백 추가
    for _ in range(4):
        doc.add_paragraph()

    # 제목
    title = doc.add_heading(get_text(locale, '통합 주간 리포트', 'Integrated Weekly Report'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 설비명
    asset_para = doc.add_paragraph()
    asset_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = asset_para.add_run(f"{asset_info.get('name', 'Unknown')} ({asset_info.get('channel', 'Main')})")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLORS['SECONDARY']

    doc.add_paragraph()
    doc.add_paragraph()

    # 리포트 기간
    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period_para.add_run(f"{get_text(locale, '리포트 기간', 'Report Period')}: {report_period}")
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS['SECONDARY']

    # 생성일시
    gen_para = doc.add_paragraph()
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_para.add_run(f"{get_text(locale, '생성일시', 'Generated')}: {generated_at}")
    run.font.size = Pt(11)
    run.font.color.rgb = COLORS['SECONDARY']

    doc.add_page_break()


def add_asset_info_section(doc, asset_info: Dict, section_num: int, locale: str):
    """설비 정보 섹션"""
    doc.add_heading(f"{section_num}. {get_text(locale, '설비 정보', 'Asset Information')}", 1)

    rows = [
        (get_text(locale, '설비명', 'Asset Name'), asset_info.get('name', '-')),
        (get_text(locale, '설비 타입', 'Asset Type'), asset_info.get('type', '-')),
        (get_text(locale, '채널', 'Channel'), asset_info.get('channel', '-')),
        (get_text(locale, '정격전압', 'Rated Voltage'), f"{asset_info.get('ratedVoltage', '-')} V"),
        (get_text(locale, '정격전류', 'Rated Current'), f"{asset_info.get('ratedCurrent', '-')} A"),
        (get_text(locale, '정격주파수', 'Rated Frequency'), f"{asset_info.get('ratedFrequency', 60)} Hz"),
    ]

    if asset_info.get('driveType'):
        drive_text = get_text(locale, '직입 기동', 'DOL') if asset_info['driveType'] == 'DOL' else \
                     get_text(locale, '인버터', 'VFD') if asset_info['driveType'] == 'VFD' else asset_info['driveType']
        rows.insert(2, (get_text(locale, '구동 방식', 'Drive Type'), drive_text))

    if asset_info.get('mac'):
        rows.append((get_text(locale, 'MAC 주소', 'MAC Address'), asset_info.get('mac', '-')))

    create_info_table(doc, rows)
    doc.add_page_break()


def add_diagnosis_section(doc, diagnosis_data: Dict, mode: str, section_num: int,
                          locale: str, temp_dir: str):
    """진단 섹션 (diagnosis / powerquality)"""
    is_pq = mode == 'powerquality'
    title = get_text(locale,
                     f"{section_num}. 전력품질 진단" if is_pq else f"{section_num}. 설비 진단",
                     f"{section_num}. Power Quality Diagnosis" if is_pq else f"{section_num}. Equipment Diagnosis")

    doc.add_heading(title, 1)

    if not diagnosis_data or not diagnosis_data.get('main'):
        p = doc.add_paragraph()
        run = p.add_run(get_text(locale, '진단 데이터가 없습니다.', 'No diagnosis data available.'))
        run.font.italic = True
        run.font.color.rgb = COLORS['SECONDARY']
        doc.add_page_break()
        return

    # 진단 시간
    if diagnosis_data.get('timestamp'):
        p = doc.add_paragraph()
        run = p.add_run(f"{get_text(locale, '진단 시간', 'Diagnosis Time')}: {diagnosis_data['timestamp']}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS['SECONDARY']

    doc.add_paragraph()

    # 바차트 생성
    main_data = diagnosis_data.get('main', [])
    if main_data:
        chart_path = os.path.join(temp_dir, f'{mode}_chart_{uuid.uuid4().hex[:8]}.png')
        logger.info(f"📊 차트 생성 시도: {chart_path}")
        try:
            create_diagnosis_bar_chart(chart_path, main_data, mode, locale)
            if os.path.exists(chart_path):
                logger.info(f"✅ 차트 파일 생성됨: {chart_path}, 크기: {os.path.getsize(chart_path)} bytes")
                doc.add_picture(chart_path, width=Inches(6.5))
                os.remove(chart_path)
            else:
                logger.error(f"❌ 차트 파일이 생성되지 않음: {chart_path}")
        except Exception as e:
            logger.error(f"❌ 차트 생성/삽입 실패: {e}")
            import traceback
            traceback.print_exc()

    doc.add_paragraph()

    # 진단 결과 테이블
    status_config = STATUS_CONFIG[mode]
    headers = [
        get_text(locale, '항목', 'Item'),
        get_text(locale, '상태', 'Status'),
        get_text(locale, '값', 'Value')
    ]

    data = []
    for item in main_data:
        status = item.get('status', 0)
        status_info = status_config.get(status, status_config[0])
        item_title = item.get('title_ko' if locale == 'ko' else 'title_en') or \
                     item.get('title') or item.get('item_name', '')

        data.append([
            item_title,
            {
                'text': status_info['text_ko' if locale == 'ko' else 'text_en'],
                'bg': status_info['bg'],
                'color': RGBColor(255, 255, 255),
                'bold': True
            },
            format_number(item.get('value'), 3)
        ])

    create_data_table(doc, headers, data, col_widths=[6, 3, 3])

    # 상세 분석 (status >= 2인 항목)
    warning_items = [item for item in main_data if item.get('status', 0) >= 2]
    detail_data = diagnosis_data.get('detail', [])

    if warning_items and detail_data:
        doc.add_paragraph()
        doc.add_heading(get_text(locale, '상세 분석', 'Detail Analysis'), 2)

        for item in warning_items:
            item_name = (item.get('item_name') or '').replace(' ', '')
            item_title = item.get('title_ko' if locale == 'ko' else 'title_en') or \
                         item.get('title') or item.get('item_name', '')
            status_info = status_config.get(item.get('status', 0), status_config[0])

            # 항목 헤더
            p = doc.add_paragraph()
            run = p.add_run(f"{item_title} - {status_info['text_ko' if locale == 'ko' else 'text_en']}")
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = status_info['color']

            # 상세 항목
            details = [d for d in detail_data
                       if (d.get('parent_name') or '').replace(' ', '') == item_name
                       and d.get('status', 0) >= 2]

            if details:
                detail_headers = [
                    get_text(locale, '세부 항목', 'Detail Item'),
                    get_text(locale, '모듈', 'Module'),
                    get_text(locale, '값', 'Value')
                ]
                detail_rows = []
                for d in details:
                    d_title = d.get('title_ko' if locale == 'ko' else 'title_en') or \
                              d.get('title') or d.get('item_name', '')
                    detail_rows.append([d_title, d.get('assembly_id', '-'), format_number(d.get('value'), 3)])

                create_data_table(doc, detail_headers, detail_rows, col_widths=[6, 3, 3], header_bg='64748B')

            # 설명
            description = item.get('description_ko' if locale == 'ko' else 'description_en') or \
                          item.get('description')
            if description:
                p = doc.add_paragraph()
                run = p.add_run(f"{description}")
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = status_info['color']

    # 트렌드 차트 (trends 데이터가 있으면 출력)
    trends_data = diagnosis_data.get('trends', [])
    logger.info(f"📈 트렌드 데이터 타입: {type(trends_data)}")

    # trends가 문자열이면 JSON 파싱
    if isinstance(trends_data, str):
        try:
            trends_data = json.loads(trends_data)
        except (json.JSONDecodeError, TypeError) as e:
            logger.error(f"❌ 트렌드 JSON 파싱 실패: {e}")
            trends_data = []

    # trends가 None이면 빈 리스트로
    if trends_data is None:
        trends_data = []

    # trends가 dict 형식이면 (item_name: [{status, timestamp}, ...]) 차트 형식으로 변환
    if isinstance(trends_data, dict) and trends_data:
        converted_trends = []
        for item_name, data_list in trends_data.items():
            if not isinstance(data_list, list) or len(data_list) == 0:
                continue

            # timestamp를 라벨로, status를 데이터로 변환
            labels = []
            values = []
            for item in data_list:
                if isinstance(item, dict):
                    ts = item.get('timestamp', '')
                    # timestamp 포맷팅 (2025-12-18T15:51:39.200000 -> 12-18 15:51)
                    if ts and 'T' in ts:
                        try:
                            date_part = ts.split('T')[0][5:]  # MM-DD
                            time_part = ts.split('T')[1][:5]  # HH:MM
                            labels.append(f"{date_part} {time_part}")
                        except:
                            labels.append(ts[:16] if len(ts) > 16 else ts)
                    else:
                        labels.append(str(ts))
                    values.append(item.get('status', 0))

            if labels and values:
                converted_trends.append({
                    'lineTitle': item_name,
                    'lineLabels': labels,
                    'lineData': [
                        {'name': 'Status', 'data': values, 'isThreshold': False}
                    ]
                })

        trends_data = converted_trends
        logger.info(f"📈 트렌드 데이터 변환 완료: {len(trends_data)}개 항목")

    # trends가 list지만 기존 형식이 아닐 수 있음
    if isinstance(trends_data, list):
        # 각 항목이 lineLabels/lineData를 갖는지 확인
        valid_trends = []
        for trend in trends_data:
            if isinstance(trend, str):
                try:
                    trend = json.loads(trend)
                except:
                    continue
            if isinstance(trend, dict) and trend.get('lineLabels') and trend.get('lineData'):
                valid_trends.append(trend)
        trends_data = valid_trends

    logger.info(f"📈 최종 트렌드 데이터 개수: {len(trends_data) if trends_data else 0}")

    if trends_data and len(trends_data) > 0:
        doc.add_paragraph()
        doc.add_heading(get_text(locale, '트렌드 차트', 'Trend Charts'), 2)

        for idx, trend in enumerate(trends_data):
            # trend가 문자열이면 파싱
            if isinstance(trend, str):
                try:
                    trend = json.loads(trend)
                except (json.JSONDecodeError, TypeError):
                    continue

            if not isinstance(trend, dict):
                continue

            if not trend.get('lineLabels') or not trend.get('lineData'):
                continue

            trend_chart_path = os.path.join(temp_dir, f'{mode}_trend_{uuid.uuid4().hex[:8]}_{idx}.png')
            result = create_trend_chart(trend_chart_path, trend, locale)

            if result and os.path.exists(trend_chart_path):
                doc.add_picture(trend_chart_path, width=Inches(6.5))
                doc.add_paragraph()
                os.remove(trend_chart_path)

    doc.add_page_break()


def add_en50160_section(doc, en50160_data: Dict, section_num: int, locale: str, temp_dir: str):
    """EN50160 분석 섹션 (시계열 차트 + 히스토그램 포함)"""
    doc.add_heading(f"{section_num}. {get_text(locale, 'EN50160 전력품질 분석', 'EN50160 Power Quality Analysis')}", 1)

    if not en50160_data:
        p = doc.add_paragraph()
        run = p.add_run(get_text(locale, 'EN50160 데이터가 없습니다.', 'No EN50160 data available.'))
        run.font.italic = True
        run.font.color.rgb = COLORS['SECONDARY']
        doc.add_page_break()
        return

    summary = en50160_data.get('summary') or {}

    # 종합 판정
    compliance = summary.get('compliance', 'N/A')
    compliance_color = '10B981' if compliance == 'PASS' else 'EF4444' if compliance == 'FAIL' else '64748B'

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, compliance_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"  {get_text(locale, '종합 판정', 'Overall Compliance')}: {compliance}  ")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # 상세 테이블 (detail_table이 있는 경우)
    detail_table = en50160_data.get('detail_table')
    if detail_table and len(detail_table) > 0:
        doc.add_heading(get_text(locale, 'EN50160 요약 테이블', 'EN50160 Summary Table'), 2)

        headers = [
            get_text(locale, '파라미터', 'Parameter'),
            'L1', 'L2', 'L3',
            get_text(locale, '판정', 'Compliance'),
            get_text(locale, '기준', 'Required Quality')
        ]

        data = []
        for row in detail_table:
            param_name = row.get('parameter_ko' if locale == 'ko' else 'parameter', '-')
            compliance_val = row.get('compliance', '-')

            # 판정 결과에 따른 스타일링
            if compliance_val in ['OK', 'PASS']:
                compliance_cell = {'text': compliance_val, 'bg': '10B981',
                                   'color': RGBColor(255, 255, 255), 'bold': True}
            elif compliance_val in ['Failed', 'FAIL']:
                compliance_cell = {'text': compliance_val, 'bg': 'EF4444',
                                   'color': RGBColor(255, 255, 255), 'bold': True}
            else:
                compliance_cell = compliance_val

            data.append([
                param_name,
                str(row.get('l1', '-')),
                str(row.get('l2', '-')),
                str(row.get('l3', '-')),
                compliance_cell,
                row.get('required', '-')
            ])

        create_data_table(doc, headers, data, col_widths=[3.5, 1.2, 1.2, 1.2, 1.8, 3.5])
        doc.add_paragraph()

    # 요약 차트
    if summary:
        chart_path = os.path.join(temp_dir, f'en50160_summary_{uuid.uuid4().hex[:8]}.png')
        create_en50160_summary_chart(chart_path, summary, locale)
        if os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(5.5))
            os.remove(chart_path)

    doc.add_paragraph()

    # ==========================================
    # 주파수 분석
    # ==========================================
    freq_data = en50160_data.get('frequency')
    if freq_data:
        doc.add_heading(get_text(locale, '주파수 분석', 'Frequency Analysis'), 2)
        stats = freq_data.get('statistics') or {}
        limits = freq_data.get('limits') or {}

        rows = [
            (get_text(locale, '공칭 주파수', 'Nominal'), f"{limits.get('nominal', 60)} Hz"),
            (get_text(locale, '최소값', 'Minimum'), f"{format_number(stats.get('min'), 3)} Hz"),
            (get_text(locale, '최대값', 'Maximum'), f"{format_number(stats.get('max'), 3)} Hz"),
            (get_text(locale, '평균값', 'Average'), f"{format_number(stats.get('avg'), 3)} Hz"),
            (get_text(locale, '99.5% 범위 내', 'In 99.5% Range'), f"{format_number(stats.get('in_range_99_5_percent'), 2)}%"),
            (get_text(locale, '판정', 'Result'), stats.get('result_99_5', 'N/A')),
        ]
        create_info_table(doc, rows)
        doc.add_paragraph()

        # 시계열 차트
        timeseries = freq_data.get('timeseries')
        if timeseries and timeseries.get('labels') and timeseries.get('data'):
            chart_path = os.path.join(temp_dir, f'freq_ts_{uuid.uuid4().hex[:8]}.png')
            title = get_text(locale, '주파수', 'Frequency') + ' (Hz)'
            create_en50160_timeseries_chart(chart_path, timeseries, title, limits, locale)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6.5))
                os.remove(chart_path)

        # 히스토그램
        histogram = freq_data.get('histogram')
        if histogram and histogram.get('bin_centers') and histogram.get('counts'):
            chart_path = os.path.join(temp_dir, f'freq_hist_{uuid.uuid4().hex[:8]}.png')
            title = get_text(locale, '주파수', 'Frequency')
            create_en50160_histogram_chart(chart_path, histogram, title, stats, limits, locale)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5.5))
                os.remove(chart_path)

        doc.add_paragraph()

    # ==========================================
    # 전압 분석
    # ==========================================
    voltage_data = en50160_data.get('voltage')
    if voltage_data and voltage_data.get('phases'):
        doc.add_heading(get_text(locale, '전압 분석', 'Voltage Analysis'), 2)
        limits = voltage_data.get('limits') or {}

        p = doc.add_paragraph()
        limit_95 = limits.get('limit_95') or {}
        run = p.add_run(f"{get_text(locale, '정격전압', 'Nominal')}: {limits.get('nominal', '-')} V, "
                        f"{get_text(locale, '95% 한계', '95% Limit')}: {limit_95.get('min', '-')} ~ "
                        f"{limit_95.get('max', '-')} V")
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['SECONDARY']

        headers = ['Phase', 'Min (V)', 'Max (V)', 'Avg (V)', '95% Range', 'Result']
        data = []
        for phase, phase_data in voltage_data['phases'].items():
            stats = (phase_data.get('statistics') if phase_data else None) or {}
            result = stats.get('result_95', 'N/A')
            data.append([
                phase,
                format_number(stats.get('min'), 1),
                format_number(stats.get('max'), 1),
                format_number(stats.get('avg'), 1),
                f"{format_number(stats.get('in_range_95_percent'), 1)}%",
                {'text': result, 'bg': '10B981' if result == 'PASS' else 'EF4444',
                 'color': RGBColor(255, 255, 255), 'bold': True}
            ])
        create_data_table(doc, headers, data, col_widths=[2, 2.5, 2.5, 2.5, 2.5, 2])
        doc.add_paragraph()

        # 각 상별 시계열 차트
        for phase, phase_data in voltage_data['phases'].items():
            if not phase_data:
                continue
            timeseries = phase_data.get('timeseries')
            if timeseries and timeseries.get('labels') and timeseries.get('data'):
                chart_path = os.path.join(temp_dir, f'volt_{phase}_ts_{uuid.uuid4().hex[:8]}.png')
                title = f"{get_text(locale, '전압', 'Voltage')} {phase} (V)"
                create_en50160_timeseries_chart(chart_path, timeseries, title, limits, locale)
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(6.5))
                    os.remove(chart_path)

            # 히스토그램
            histogram = phase_data.get('histogram')
            if histogram and histogram.get('bin_centers') and histogram.get('counts'):
                stats = phase_data.get('statistics') or {}
                chart_path = os.path.join(temp_dir, f'volt_{phase}_hist_{uuid.uuid4().hex[:8]}.png')
                title = f"{get_text(locale, '전압', 'Voltage')} {phase}"
                create_en50160_histogram_chart(chart_path, histogram, title, stats, limits, locale)
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(5.5))
                    os.remove(chart_path)
            doc.add_paragraph()

    # ==========================================
    # THD 분석
    # ==========================================
    thd_data = en50160_data.get('thd')
    if thd_data and thd_data.get('phases'):
        doc.add_heading(get_text(locale, 'THD 분석', 'THD Analysis'), 2)

        thd_limits = thd_data.get('limits') or {}
        p = doc.add_paragraph()
        run = p.add_run(f"{get_text(locale, '95% 한계', '95% Limit')}: {thd_limits.get('limit_95', 8)}%")
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['SECONDARY']

        headers = ['Phase', 'Max (%)', 'Avg (%)', '95th (%)', '95% Range', 'Result']
        data = []
        for phase, phase_data in thd_data['phases'].items():
            stats = (phase_data.get('statistics') if phase_data else None) or {}
            result = stats.get('result', 'N/A')
            data.append([
                phase,
                format_number(stats.get('max'), 2),
                format_number(stats.get('avg'), 2),
                format_number(stats.get('percentile_95'), 2),
                f"{format_number(stats.get('in_range_95_percent'), 1)}%",
                {'text': result, 'bg': '10B981' if result == 'PASS' else 'EF4444',
                 'color': RGBColor(255, 255, 255), 'bold': True}
            ])
        create_data_table(doc, headers, data, col_widths=[2, 2.5, 2.5, 2.5, 2.5, 2])
        doc.add_paragraph()

        # 각 상별 차트 (첫번째 상만 출력)
        for phase, phase_data in list(thd_data['phases'].items())[:1]:
            if not phase_data:
                continue
            timeseries = phase_data.get('timeseries')
            if timeseries and timeseries.get('labels') and timeseries.get('data'):
                chart_path = os.path.join(temp_dir, f'thd_{phase}_ts_{uuid.uuid4().hex[:8]}.png')
                title = f"THD {phase} (%)"
                create_en50160_timeseries_chart(chart_path, timeseries, title, thd_limits, locale)
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(6.5))
                    os.remove(chart_path)

            # 히스토그램
            histogram = phase_data.get('histogram')
            if histogram and histogram.get('bin_centers') and histogram.get('counts'):
                chart_path = os.path.join(temp_dir, f'thd_{phase}_hist_{uuid.uuid4().hex[:8]}.png')
                title = f"THD {phase}"
                stats = phase_data.get('statistics') or {}
                create_en50160_histogram_chart(chart_path, histogram, title, stats, thd_limits, locale)
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(5.5))
                    os.remove(chart_path)

    # ==========================================
    # 불평형 분석
    # ==========================================
    unbal_data = en50160_data.get('unbalance')
    if unbal_data:
        doc.add_heading(get_text(locale, '전압 불평형 분석', 'Voltage Unbalance Analysis'), 2)
        stats = unbal_data.get('statistics') or {}
        unbal_limits = unbal_data.get('limits') or {}

        rows = [
            (get_text(locale, '95% 한계', '95% Limit'), f"{unbal_limits.get('limit_95', 2)}%"),
            (get_text(locale, '최대값', 'Maximum'), f"{format_number(stats.get('max'), 3)}%"),
            (get_text(locale, '평균값', 'Average'), f"{format_number(stats.get('avg'), 3)}%"),
            (get_text(locale, '95번째 백분위', '95th Percentile'), f"{format_number(stats.get('percentile_95'), 3)}%"),
            (get_text(locale, '95% 범위 내', 'In 95% Range'), f"{format_number(stats.get('in_range_95_percent'), 2)}%"),
            (get_text(locale, '판정', 'Result'), stats.get('result', 'N/A')),
        ]
        create_info_table(doc, rows)
        doc.add_paragraph()

        # 시계열 차트
        timeseries = unbal_data.get('timeseries')
        if timeseries and timeseries.get('labels') and timeseries.get('data'):
            chart_path = os.path.join(temp_dir, f'unbal_ts_{uuid.uuid4().hex[:8]}.png')
            title = get_text(locale, '전압 불평형', 'Voltage Unbalance') + ' (%)'
            create_en50160_timeseries_chart(chart_path, timeseries, title, unbal_limits, locale)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6.5))
                os.remove(chart_path)

        # 히스토그램
        histogram = unbal_data.get('histogram')
        if histogram and histogram.get('bin_centers') and histogram.get('counts'):
            chart_path = os.path.join(temp_dir, f'unbal_hist_{uuid.uuid4().hex[:8]}.png')
            title = get_text(locale, '전압 불평형', 'Voltage Unbalance')
            create_en50160_histogram_chart(chart_path, histogram, title, stats, unbal_limits, locale)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5.5))
                os.remove(chart_path)

    # ==========================================
    # 플리커 분석
    # ==========================================
    flicker_data = en50160_data.get('flicker')
    if flicker_data and flicker_data.get('plt'):
        doc.add_heading(get_text(locale, '플리커 분석 (Plt)', 'Flicker Analysis (Plt)'), 2)

        headers = ['Phase', 'Max', 'Avg', '95th', '95% Range', 'Result']
        data = []
        for phase, phase_data in flicker_data['plt'].items():
            stats = (phase_data.get('statistics') if phase_data else None) or {}
            result = stats.get('result', 'N/A')
            data.append([
                phase,
                format_number(stats.get('max'), 3),
                format_number(stats.get('avg'), 3),
                format_number(stats.get('percentile_95'), 3),
                f"{format_number(stats.get('in_range_95_percent'), 1)}%",
                {'text': result, 'bg': '10B981' if result == 'PASS' else 'EF4444',
                 'color': RGBColor(255, 255, 255), 'bold': True}
            ])
        create_data_table(doc, headers, data, col_widths=[2, 2.5, 2.5, 2.5, 2.5, 2])
        doc.add_paragraph()

        # 첫번째 상 차트만
        for phase, phase_data in list(flicker_data['plt'].items())[:1]:
            if not phase_data:
                continue
            timeseries = phase_data.get('timeseries')
            if timeseries and timeseries.get('labels') and timeseries.get('data'):
                chart_path = os.path.join(temp_dir, f'plt_{phase}_ts_{uuid.uuid4().hex[:8]}.png')
                title = f"Plt {phase}"
                flicker_limits = flicker_data.get('limits') or {}
                create_en50160_timeseries_chart(chart_path, timeseries, title, flicker_limits, locale)
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(6.5))
                    os.remove(chart_path)

    # ==========================================
    # 고조파 분석
    # ==========================================
    harmonics_data = en50160_data.get('harmonics')
    if harmonics_data and harmonics_data.get('phases'):
        doc.add_heading(get_text(locale, '고조파 분석', 'Harmonics Analysis'), 2)

        harm_limits = harmonics_data.get('limits') or {}

        # 각 상별 테이블
        for phase, phase_data in harmonics_data['phases'].items():
            if not phase_data:
                continue

            p = doc.add_paragraph()
            run = p.add_run(f"{phase}")
            run.font.bold = True
            run.font.size = Pt(11)

            # 고조파 테이블
            headers = [
                get_text(locale, '고조파', 'Harmonic'),
                get_text(locale, '최대 (%)', 'Max (%)'),
                get_text(locale, '평균 (%)', 'Avg (%)'),
                get_text(locale, '한계 (%)', 'Limit (%)'),
                get_text(locale, '95% 범위', '95% Range'),
                get_text(locale, '판정', 'Result')
            ]
            data = []

            for h_key in ['h2', 'h3', 'h4', 'h5', 'h6', 'h7', 'h8', 'h9', 'h10', 'h11', 'h12', 'h13', 'h14', 'h15', 'h16', 'h17', 'h18', 'h19', 'h20', 'h21', 'h22', 'h23', 'h24', 'h25']:
                h_data = phase_data.get(h_key)
                if not h_data:
                    continue

                h_num = h_key.upper()
                result = h_data.get('result', 'N/A')
                data.append([
                    h_num,
                    format_number(h_data.get('max'), 2),
                    format_number(h_data.get('avg'), 2),
                    format_number(h_data.get('limit'), 1),
                    f"{format_number(h_data.get('in_range_95_percent'), 1)}%",
                    {'text': result, 'bg': '10B981' if result == 'PASS' else 'EF4444',
                     'color': RGBColor(255, 255, 255), 'bold': True}
                ])

            if data:
                create_data_table(doc, headers, data, col_widths=[1.5, 2, 2, 2, 2.5, 1.5])
            doc.add_paragraph()

            # 첫번째 상만 히스토그램
            if phase == 'L1':
                histogram = phase_data.get('histogram')
                if histogram and histogram.get('bin_centers') and histogram.get('counts'):
                    chart_path = os.path.join(temp_dir, f'harm_{phase}_hist_{uuid.uuid4().hex[:8]}.png')
                    title = get_text(locale, '고조파 THD', 'Harmonics THD') + f" {phase}"
                    create_en50160_histogram_chart(chart_path, histogram, title, {}, {}, locale)
                    if os.path.exists(chart_path):
                        doc.add_picture(chart_path, width=Inches(5.5))
                        os.remove(chart_path)

    doc.add_page_break()


def add_energy_section(doc, energy_data: Dict, section_num: int, locale: str, temp_dir: str):
    """전력량 분석 섹션 (시간별 → 일별 → 월별 순서)"""
    doc.add_heading(f"{section_num}. {get_text(locale, '전력량 분석', 'Energy Analysis')}", 1)

    if not energy_data:
        p = doc.add_paragraph()
        run = p.add_run(get_text(locale, '전력량 데이터가 없습니다.', 'No energy data available.'))
        run.font.italic = True
        run.font.color.rgb = COLORS['SECONDARY']
        return

    # ==========================================
    # 1. 시간대별 사용 패턴
    # ==========================================
    hourly_pattern = energy_data.get('hourlyPattern', {})
    hourly_data = energy_data.get('hourly', [])

    if hourly_pattern or hourly_data:
        doc.add_heading(get_text(locale, '시간대별 사용 패턴', 'Hourly Usage Pattern'), 2)

        if hourly_pattern:
            peak = hourly_pattern.get('peakHour', {})
            off_peak = hourly_pattern.get('offPeakHour', {})

            ratio = '-'
            if off_peak.get('avgValue', 0) > 0:
                ratio = f"{peak.get('avgValue', 0) / off_peak['avgValue']:.2f}x"

            rows = [
                (get_text(locale, '피크 시간대', 'Peak Hour'),
                 f"{peak.get('hour', '-')}:00 ({format_number(peak.get('avgValue'), 2)} kWh)"),
                (get_text(locale, '오프피크 시간대', 'Off-Peak Hour'),
                 f"{off_peak.get('hour', '-')}:00 ({format_number(off_peak.get('avgValue'), 2)} kWh)"),
                (get_text(locale, '피크/오프피크 비율', 'Peak/Off-Peak Ratio'), ratio),
            ]
            create_info_table(doc, rows)

        # 시간별 차트
        if hourly_data:
            doc.add_paragraph()
            chart_path = os.path.join(temp_dir, f'energy_hourly_{uuid.uuid4().hex[:8]}.png')
            create_hourly_energy_chart(chart_path, hourly_data, locale)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6))
                os.remove(chart_path)

        doc.add_paragraph()

    # ==========================================
    # 2. 일별 전력량
    # ==========================================
    daily_data = energy_data.get('daily', [])

    if daily_data:
        # 통계 계산
        values = [d.get('value', 0) for d in daily_data]
        total = sum(values)
        avg = total / len(values) if values else 0
        max_val = max(values) if values else 0
        min_val = min(values) if values else 0
        max_day = next((d for d in daily_data if d.get('value') == max_val), {})
        min_day = next((d for d in daily_data if d.get('value') == min_val), {})

        doc.add_heading(get_text(locale, '일별 전력량 통계', 'Daily Energy Statistics'), 2)

        rows = [
            (get_text(locale, '총 전력량', 'Total Energy'), f"{format_number(total, 2)} kWh"),
            (get_text(locale, '일 평균', 'Daily Average'), f"{format_number(avg, 2)} kWh"),
            (get_text(locale, '최대 사용일', 'Maximum Day'),
             f"{max_day.get('date', '-')} ({format_number(max_val, 2)} kWh)"),
            (get_text(locale, '최소 사용일', 'Minimum Day'),
             f"{min_day.get('date', '-')} ({format_number(min_val, 2)} kWh)"),
        ]
        create_info_table(doc, rows)

        doc.add_paragraph()

        # 일별 차트
        chart_path = os.path.join(temp_dir, f'energy_chart_{uuid.uuid4().hex[:8]}.png')
        create_energy_chart(chart_path, daily_data, locale)
        if os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(6))
            os.remove(chart_path)

        doc.add_paragraph()

        # 일별 데이터 테이블
        doc.add_heading(get_text(locale, '일별 상세', 'Daily Details'), 2)

        headers = [
            get_text(locale, '날짜', 'Date'),
            get_text(locale, '전력량 (kWh)', 'Energy (kWh)'),
            get_text(locale, '비율', 'Ratio')
        ]
        data = []
        for d in daily_data[:14]:  # 최대 14일
            ratio = (d.get('value', 0) / max_val * 100) if max_val > 0 else 0
            data.append([
                d.get('date', '-'),
                format_number(d.get('value'), 2),
                f"{format_number(ratio, 1)}%"
            ])
        create_data_table(doc, headers, data, col_widths=[4, 4, 4])

        doc.add_paragraph()

    # ==========================================
    # 3. 월별 전력량
    # ==========================================
    monthly_data = energy_data.get('monthly', [])
    if monthly_data:
        doc.add_heading(get_text(locale, '월별 전력량', 'Monthly Energy'), 2)

        # 월별 차트
        chart_path = os.path.join(temp_dir, f'energy_monthly_{uuid.uuid4().hex[:8]}.png')
        create_monthly_energy_chart(chart_path, monthly_data, locale)
        if os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(6))
            os.remove(chart_path)


def create_itic_curve_chart(chart_path: str, itic_events: List[Dict], locale: str = 'en') -> str:
    """
    ITIC Curve 차트 생성

    Args:
        chart_path: 저장 경로
        itic_events: ITIC 이벤트 리스트 [{duration, voltage_pct, event_type}, ...]
        locale: 언어

    Returns:
        저장된 파일 경로
    """
    font_prop = setup_korean_font()

    fig, ax = plt.subplots(figsize=(10, 7))

    # ITIC Curve 영역 정의
    # Prohibited Region (하단)
    prohibited_lower_x = [0.001, 0.001, 0.02, 0.02, 0.5, 0.5, 10, 10, 0.001]
    prohibited_lower_y = [0, 70, 70, 40, 40, 70, 70, 0, 0]
    ax.fill(prohibited_lower_x, prohibited_lower_y,
            color='#fee2e2', alpha=0.7, label=get_text(locale, '금지 영역', 'Prohibited Region'))

    # No Damage Region (중간)
    no_damage_x = [0.001, 0.001, 0.003, 0.003, 0.02, 0.5, 10, 10, 0.5, 0.02, 0.003, 0.001]
    no_damage_y = [70, 110, 110, 120, 120, 120, 110, 70, 70, 40, 70, 70]
    ax.fill(no_damage_x, no_damage_y,
            color='#d1fae5', alpha=0.7, label=get_text(locale, '정상 영역', 'No Damage Region'))

    # Prohibited Region (상단)
    prohibited_upper_x = [0.001, 0.001, 0.003, 0.003, 0.5, 10, 10, 0.5, 0.003, 0.001]
    prohibited_upper_y = [110, 200, 200, 140, 140, 120, 110, 120, 120, 110]
    ax.fill(prohibited_upper_x, prohibited_upper_y,
            color='#fee2e2', alpha=0.7)

    # ITIC Curve 라인
    ax.plot([0.001, 0.003, 0.003, 0.5, 10], [110, 110, 120, 120, 110],
            'r-', linewidth=2, label='ITIC Upper Limit')
    ax.plot([0.001, 0.02, 0.02, 0.5, 10], [70, 70, 40, 70, 70],
            'r-', linewidth=2, label='ITIC Lower Limit')

    # 이벤트 플롯
    if itic_events:
        sag_events = [e for e in itic_events if e.get('event_type') == 'SAG']
        swell_events = [e for e in itic_events if e.get('event_type') == 'SWELL']

        if sag_events:
            sag_x = [e['duration'] for e in sag_events]
            sag_y = [e['voltage_pct'] for e in sag_events]
            ax.scatter(sag_x, sag_y, c='orange', s=60, alpha=0.8,
                      edgecolors='black', linewidth=1, zorder=5,
                      label=f'SAG ({len(sag_events)})')

        if swell_events:
            swell_x = [e['duration'] for e in swell_events]
            swell_y = [e['voltage_pct'] for e in swell_events]
            ax.scatter(swell_x, swell_y, c='green', s=60, alpha=0.8,
                      edgecolors='black', linewidth=1, zorder=5,
                      label=f'SWELL ({len(swell_events)})')

    ax.set_xscale('log')
    ax.set_xlim(0.001, 10)
    ax.set_ylim(0, 200)

    xlabel = get_text(locale, '지속 시간 (초)', 'Duration (s)')
    ylabel = get_text(locale, '전압 (%)', 'Voltage (%)')
    title = get_text(locale, 'ITIC Curve - 전력품질 이벤트', 'ITIC Curve - Power Quality Events')

    if font_prop:
        ax.set_xlabel(xlabel, fontsize=11, fontweight='bold', fontproperties=font_prop)
        ax.set_ylabel(ylabel, fontsize=11, fontweight='bold', fontproperties=font_prop)
        ax.set_title(title, fontsize=13, fontweight='bold', pad=15, fontproperties=font_prop)
    else:
        ax.set_xlabel(xlabel, fontsize=11, fontweight='bold')
        ax.set_ylabel(ylabel, fontsize=11, fontweight='bold')
        ax.set_title(title, fontsize=13, fontweight='bold', pad=15)

    ax.grid(True, which='both', alpha=0.3, linestyle='--')
    ax.legend(loc='upper right', fontsize=9)

    # 100% 기준선
    ax.axhline(y=100, color='blue', linestyle=':', alpha=0.5, linewidth=1)
    ax.text(0.002, 102, '100%', fontsize=8, color='blue', alpha=0.7)

    plt.tight_layout()
    plt.savefig(chart_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()

    return chart_path


def add_itic_section(doc, itic_events: List[Dict], section_num: int, locale: str, temp_dir: str):
    """ITIC Curve 분석 섹션"""
    doc.add_heading(f"{section_num}. {get_text(locale, 'ITIC Curve 분석', 'ITIC Curve Analysis')}", 1)

    if not itic_events:
        p = doc.add_paragraph()
        run = p.add_run(get_text(locale, 'ITIC 이벤트가 없습니다.', 'No ITIC events detected.'))
        run.font.italic = True
        run.font.color.rgb = COLORS['SECONDARY']
        doc.add_page_break()
        return

    # 이벤트 요약
    sag_count = len([e for e in itic_events if e.get('event_type') == 'SAG'])
    swell_count = len([e for e in itic_events if e.get('event_type') == 'SWELL'])
    total_count = len(itic_events)

    p = doc.add_paragraph()
    run = p.add_run(get_text(locale,
        f"총 {total_count}개의 전압 이벤트가 감지되었습니다.",
        f"Total {total_count} voltage events detected."))
    run.font.size = Pt(11)

    doc.add_paragraph()

    # 요약 테이블
    rows = [
        (get_text(locale, '전체 이벤트', 'Total Events'), str(total_count)),
        (get_text(locale, '전압 순간저하 (SAG)', 'Voltage Sag (SAG)'), str(sag_count)),
        (get_text(locale, '전압 순간상승 (SWELL)', 'Voltage Swell (SWELL)'), str(swell_count)),
    ]
    create_info_table(doc, rows)

    doc.add_paragraph()

    # ITIC Curve 차트
    chart_path = os.path.join(temp_dir, f'itic_chart_{uuid.uuid4().hex[:8]}.png')
    create_itic_curve_chart(chart_path, itic_events, locale)
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6))
        os.remove(chart_path)

    doc.add_paragraph()

    # 이벤트 상세 테이블 (최근 20개만)
    if len(itic_events) > 0:
        doc.add_heading(get_text(locale, '이벤트 상세', 'Event Details'), 2)

        headers = [
            get_text(locale, '발생 시간', 'Event Time'),
            get_text(locale, '유형', 'Type'),
            get_text(locale, '지속 시간', 'Duration'),
            get_text(locale, '전압 (%)', 'Voltage (%)')
        ]

        data = []
        for event in itic_events[:20]:  # 최대 20개
            event_type = event.get('event_type', '-')
            type_color = COLORS['WARNING'] if event_type == 'SAG' else COLORS['SUCCESS']

            data.append([
                event.get('event_time', '-') or '-',
                {
                    'text': event_type,
                    'bg': 'F97316' if event_type == 'SAG' else '10B981',
                    'color': RGBColor(255, 255, 255),
                    'bold': True
                },
                f"{event.get('duration', 0):.3f} s",
                f"{event.get('voltage_pct', 0):.1f}%"
            ])

        create_data_table(doc, headers, data, col_widths=[4, 2, 3, 3])

        if len(itic_events) > 20:
            p = doc.add_paragraph()
            run = p.add_run(get_text(locale,
                f'... 외 {len(itic_events) - 20}개 이벤트',
                f'... and {len(itic_events) - 20} more events'))
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = COLORS['SECONDARY']

    doc.add_page_break()


# ============================================
# 메인 리포트 생성 함수
# ============================================
def generate_weekly_report(
    asset_info: Dict,
    diagnosis_data: Optional[Dict],
    powerquality_data: Optional[Dict],
    en50160_data: Optional[Dict],
    energy_data: Optional[Dict],
    itic_events: Optional[List[Dict]] = None,
    report_period: str = '',
    generated_at: str = '',
    locale: str = 'en',
    output_path: str = None
) -> str:
    """
    통합 주간 리포트 생성

    Args:
        asset_info: 설비 정보
        diagnosis_data: 설비 진단 데이터
        powerquality_data: 전력품질 진단 데이터
        en50160_data: EN50160 데이터
        energy_data: 전력량 데이터
        itic_events: ITIC 이벤트 데이터
        report_period: 리포트 기간
        generated_at: 생성 시간
        locale: 언어 (ko/en)
        output_path: 출력 파일 경로

    Returns:
        생성된 파일 경로
    """
    temp_dir = tempfile.gettempdir()

    if not output_path:
        output_path = os.path.join(temp_dir, f'weekly_report_{uuid.uuid4().hex[:8]}.docx')

    doc = Document()

    # 문서 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # 1. 표지
    add_cover_page(doc, asset_info, report_period, generated_at, locale)

    section_num = 1

    # 2. 설비 정보
    add_asset_info_section(doc, asset_info, section_num, locale)
    section_num += 1

    # 3. 설비 진단
    if diagnosis_data:
        add_diagnosis_section(doc, diagnosis_data, 'diagnosis', section_num, locale, temp_dir)
        section_num += 1

    # 4. 전력품질 진단
    if powerquality_data:
        add_diagnosis_section(doc, powerquality_data, 'powerquality', section_num, locale, temp_dir)
        section_num += 1

    # 5. EN50160 분석
    if en50160_data:
        add_en50160_section(doc, en50160_data, section_num, locale, temp_dir)
        section_num += 1

    # 6. ITIC Curve 분석
    if itic_events:
        add_itic_section(doc, itic_events, section_num, locale, temp_dir)
        section_num += 1

    # 7. 전력량 분석
    if energy_data:
        add_energy_section(doc, energy_data, section_num, locale, temp_dir)

    doc.save(output_path)
    logger.info(f"✅ 통합 주간 리포트 생성 완료: {output_path}")

    return output_path
